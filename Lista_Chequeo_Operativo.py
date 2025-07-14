import streamlit as st
import os
import json
import tempfile
import io
import pandas as pd
from datetime import date
import streamlit as st
import requests
from io import BytesIO
import sys
from docx import Document

# Estilo para tonos azules y logo
st.markdown(
    """
    <style>
    /* Cambiar el color de los radio buttons */
    div[data-baseweb="radio"] label {
        color: #005A9C !important; /* Azul oscuro */
    }
    
    /* Cambiar el color del fondo cuando se selecciona */
    div[data-baseweb="radio"] input:checked + div {
        background-color: #0078D7 !important; /* Azul claro */
        color: white !important;
    }

    /* Cambiar color de fondo de la barra lateral */
    section[data-testid="stSidebar"] {
        background-color: #D9EAF7 !important; /* Azul muy claro */
    }

    /* Cambiar color de los botones */
    button {
        background-color: #0078D7 !important; /* Azul */
        color: white !important;
    }
    button:hover {
        background-color: #005A9C !important; /* Azul más oscuro */
    }
    
    /* Cambiar el color del punto del radio button sin afectar el fondo del texto */
    [role="radiogroup"] > label > div:first-child {
    background-color: #005A9C !important;  /* Azul oscuro */
    border-color: #005A9C !important;  /* Azul oscuro */
    }

    /* Mantener el fondo del texto transparente */
    [role="radiogroup"] > label > div:nth-child(2) {
    background-color: transparent !important;
    }
             
    </style>
    """,
    unsafe_allow_html=True
)

# Agregar el logo
logo_path = "logo2.png"  # Cambia la ruta al archivo de tu logo
st.sidebar.image(logo_path, use_container_width=True)

# Tu código principal
st.title("INSTRUMENTO DE VERIFICACIÓN DE SERVICIO FARMACÉUTICO")

# Diccionario de operadores y NITs
Operadores = {
    "SELECCIONAR": "---------------",
    "COHAN": "890985122",
    "GENHOSPI": "900331412",
    "INSUMEDIC": "900716566",
    "MEDICINA Y TECNOLOGIA EN SALUD MYT": "900057926",
    "SUMINISTROS Y DOTACIONES": "802000608",
    "DISFARMA": "900580962",
    "CRUZ VERDE": "800149695",
    "DISCOLMETS": "828002423",
    "HUMSALUD": "900994906",
    "ETICOS": "892300678",
    "JEZA": "901048992",
    "MEDISFARMA": "901148499",
    "DISTRIMEVD": "901433283",
    "CYA": "860029216",
    "AUDIFARMA": "816001182",
    "FARMACIA INST": "900285194",
    "INHAPRO": "901200716",
    "PHARMASAN": "900249425",
    "CLINICA EL CARMEN": "800149384",
    "VALENTECH PHARMA COLOMBIA SAS": "900677118",
    "LIGA CONTRA EL CANCER SECCIONAL RISARALDA": "891408586",
    "CENTRO ONCOLOGICO DE ANTIOQUIA SAS": "900236850",
    "UNIDAD HEMATO ONCOLOGICA Y DE RADIOTERAPIA DEL CARIBE": "900095504",
    "CLINICA LA ERMITA DE CARTAGENA": "900491883",
    "UNION TEMPORAL ONCOLOGICA ISNOOS SN": "901776252",
    "SANOFI AVENTIS DE COLOMBIA SA": "830010337",
    "NEXT PHARMA COLOMBIA SAS": "900382525",
    "EXELTIS SAS": "900716452",
    "FORPRESALUD": "804008792",
    "MACROMED": "830107855",
    "MEDYTEC": "900057926",
    "OFFIMEDICAS": "900098550",
    "ETTICOS": "892300678",
    "MULTICARE": "901671387"
}

# Inicializar variables de sesión
if "responses" not in st.session_state:
    st.session_state["responses"] = {}
if "consecutivo" not in st.session_state:
    st.session_state["consecutivo"] = 1
if "form" not in st.session_state:
    st.session_state["form"] = {}
if "load_form" not in st.session_state:
    st.session_state["load_form"]=False   
if 'Datos Generales' not in st.session_state:
    st.session_state['Datos Generales'] = {}

# Ruta de almacenamiento de formularios
folder_path = "formularios_guardados"
os.makedirs(folder_path, exist_ok=True)

def convertir_fecha(obj):
    if isinstance(obj, date):
        return obj.isoformat()  # Convierte a cadena 'YYYY-MM-DD'
    raise TypeError("Tipo no serializable")


# Función para guardar el estado actual del formulario
def guardar_estado():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = date.today().isoformat()
    consecutivo = f"{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}.json"
    file_path = os.path.join(folder_path, filename)
    

    # Crear contenido JSON
    data = {
        "form": st.session_state["form"],
        "responses": st.session_state["responses"],
        "consecutivo": st.session_state["consecutivo"],
        "load_form": st.session_state["load_form"]
    }
    
    
    # Guardar en un archivo temporal para permitir la descarga
    with open(file_path, "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=4, default=convertir_fecha)

    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")

    # Botón para descargar
    with open(file_path, "r",encoding="utf-8") as file:
        st.download_button(
            label="Descargar formulario",
            data=file.read(),
            file_name=filename,
            mime="application/json",
        )

# Función para cargar formulario desde un archivo JSON basado en un consecutivo
def cargar_formulario_por_consecutivo(consecutivo_input):
    # Ahora la función acepta un argumento
    archivo_a_cargar = f"Formulario_{str(consecutivo_input)}.json"
    file_path = os.path.join(folder_path, archivo_a_cargar)

    if os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Actualizar session_state con los datos cargados
            st.session_state["Datos Generales"] = data.get("Datos Generales",{})
            st.session_state["form"] = data.get("form", {})
            st.session_state["responses"] = data.get("responses", {})
            st.session_state["consecutivo"] = data.get("consecutivo", 1)
            st.session_state["form"]["observacion"] = st.session_state["form"].get("observacion", "")
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Variables adicionales de datos generales
            operador = st.session_state["form"].get("Operador", "SELECCIONAR")
            nit_operador = operadores.get(operador, "DESCONOCIDO")
            farmacia_seleccionada = st.session_state["form"].get("farmacia_seleccionada", "NO SELECCIONADA")
            Nit_sucursal = farmacias_filtradas.get("Nit_sucursal:", datos_farmacia.get("COD. SUC", "DESCONOCIDO"))
            ciudad = st.session_state["form"].get("ciudad", "CIUDAD NO ESPECIFICADA")
            direccion = st.session_state["form"].get("direccion", "DIRECCIÓN NO ESPECIFICADA")
            telefono = st.session_state["form"].get("telefono", "TELEFONO NO ESPECIFICADO")

            # Mostrar los campos en el formulario
            st.text_input("Operador", value=operador)
            st.text_input("Nit del Operador", value=nit_operador)
            st.text_input("Farmacia Seleccionada", value=farmacia_seleccionada)
            st.text_input("Nit de la Sucursal", value=Nit_sucursal)
            st.text_input("Ciudad", value=ciudad)
            st.text_input("Dirección", value=direccion)
            st.text_input("Teléfono", value=telefono)

            # Mostrar otros campos del formulario
            st.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
            st.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
            st.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
            st.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
            st.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
            st.text_input("Fecha de Auditoría", value=st.session_state["form"].get("Fecha de Auditoria", ""))

            # Mostrar la observación en un área de texto
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Escribir los datos generales y respuestas para la visualización
            Contenido = ""
            for grupo, preguntas in st.session_state["responses"].items():
                for pregunta, respuesta in preguntas.items():
                    Contenido += (
                        f"{operador}|{nit_operador}|"
                        f"{farmacia_seleccionada}|"
                        f"{Nit_sucursal}|"
                        f"{ciudad}|"
                        f"{direccion}|"
                        f"{telefono}|"
                        f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                        f"{st.session_state['form'].get('Representante legal', '')}|"
                        f"{st.session_state['form'].get('Director técnico', '')}|"
                        f"{st.session_state['form'].get('Auditor 1', '')}|"
                        f"{st.session_state['form'].get('Auditor 2', '')}|"
                        f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    )
            st.write("Contenido generado:")
            st.write(Contenido)

            st.success(f"Formulario {archivo_a_cargar} cargado correctamente.")
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
    else:
        st.error(f"No se encontró un archivo con el consecutivo {consecutivo_input}.")


# Función para finalizar y generar el archivo TXT
def finalizar_formulario():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    nit_operador = Operadores.get(Operador, "DESCONOCIDO")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = "2025"
    consecutivo = f"{fecha_auditoria[:4]}_{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}.txt"
    filename_word = f"Acta_Seguimiento_{farmacia_seleccionada}.docx"
    folder_path = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/Acta_Seguimiento.docx"

    # Descargar el archivo .docx desde GitHub
    response = requests.get(folder_path)
    if response.status_code == 200:
        # Cargar el archivo .docx directamente desde la respuesta
        doc = Document(BytesIO(response.content))
    else:
        st.error(f"Error al descargar el archivo: {response.status_code}")
        return

    # Continuar con el resto de tu lógica
    file_path = os.path.join(folder_path, filename)
    file_path_word = os.path.join(folder_path, filename_word)

  
     # Escribir cabecera
    Contenido = "Operador|Nit operador|Nombre de la droguería o farmacia|Nit_sucursal|Ciudad|Dirección|Teléfono|"
    Contenido += "Nivel y Tipo de servicio farmacéutico|Representante legal|Director técnico|"
    Contenido += "Auditor 1|Auditor 2|Fecha de Auditoria|Grupo de Pregunta|Respuesta|Valor|Observacion|Subgrupo de pregunta\n"

    Operador = st.session_state["form"].get("Operador","")
    nit_operador = st.session_state["form"].get("Nit_operador", "DESCONOCIDO")
    nombre_farmacia = st.session_state["form"].get("Nombre")
    nivel_servicio = st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", "")
    representante_legal = st.session_state["form"].get("Representante legal", "")
    director_tecnico = st.session_state["form"].get("Director técnico", "")
    auditor_2 = st.session_state["form"].get("Auditor 2", "")
    Observacion = st.session_state["form"].get("Observacion", "")
    tipo_drogueria = st.session_state["form"].get("Tipo de Droguería", "")
    
        
     # Escribir datos generales y respuestas
    for grupo, preguntas in st.session_state["responses"].items():
            for pregunta, respuesta in preguntas.items():
                Contenido +=(
                    f"{Operador}|{nit_operador}|"
                    f"{farmacia_seleccionada}|"
                    f"{Nit_sucursal}|"
                    f"{ciudad}|"
                    f"{direccion}|"
                    f"{telefono}|"
                    f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                    f"{st.session_state['form'].get('Representante legal', '')}|"
                    f"{st.session_state['form'].get('Director técnico', '')}|"
                    f"{st.session_state['form'].get('Auditor 1', '')}|"
                    f"{st.session_state['form'].get('Auditor 2', '')}|"
                    f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    f"{grupo}|{respuesta['respuesta']}|{respuesta['valor']}|"
                    f"{Observacion}|"
                    f"{pregunta}\n"
                )
    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")    
            
    # Convertir el contenido a bytes
    Contenido_bytes = Contenido.encode("utf-8")
    
    #boton de descarga
    st.download_button(label="Descargar archivo txt",
                   data=Contenido,
                   file_name=filename,
                   mime="text/plain"
                   )

    # Crear archivo txt
    os.makedirs(folder_path, exist_ok=True)
    with open(file_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(Contenido) 

    # Rellenar los campos de la plantilla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("OPERADOR", Operador)
                if "NIT_OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_OPERADOR", nit_operador)
                if "NOMBRE_FARMACIA" in cell.text.strip():
                    cell.text = cell.text.replace("NOMBRE_FARMACIA", farmacia_seleccionada)
                if "NIT_SUCURSAL" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_SUCURSAL", Nit_sucursal)
                if "CIUDAD" in cell.text.strip():
                    cell.text = cell.text.replace("CIUDAD", ciudad)
                if "DIRECCION" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECCION", direccion)
                if "TELEFONO" in cell.text.strip():
                    cell.text = cell.text.replace("TELEFONO", telefono)
                if "NIVEL_SERVICIO" in cell.text.strip():
                    cell.text = cell.text.replace("NIVEL_SERVICIO", nivel_servicio)
                if "REPRESENTANTE_LEGAL" in cell.text.strip():
                    cell.text = cell.text.replace("REPRESENTANTE_LEGAL", representante_legal)
                if "DIRECTOR_TECNICO" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECTOR_TECNICO", director_tecnico)
                if "AUDITOR_1" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_1", Auditor)
                if "AUDITOR_2" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_2", auditor_2)
                if "FECHA_AUDITORIA" in cell.text.strip():
                    cell.text = cell.text.replace("FECHA_AUDITORIA", fecha_auditoria)
                if "TIPO_DROGERIA" in cell.text.strip():
                    cell.text = cell.text.replace("TIPO_DROGERIA", tipo_drogueria)    
                    
    # Crear un string que contendrá todo el contenido a insertar en el documento
    contenido_respuestas = ""
    contenido_observaciones = ""
    
    # Recorrer los grupos y subgrupos de preguntas
    observaciones_contador = 1
    for grupo, subgrupos in st.session_state["responses"].items():
        # Primero agregamos el nombre del grupo, que se repetirá solo una vez
        contenido_respuestas += f"**{grupo}**\n"  # Agregar el grupo en negrita
        
        for subgrupo, respuesta_data in subgrupos.items():
            # Formatear el texto para cada grupo y subgrupo
            contenido_respuestas += f"{subgrupo}\n\n"

            # Respuesta
            respuesta = respuesta_data.get('respuesta', 'sin respuesta')
            contenido_respuestas += f"**Respuesta:** {respuesta}\n"
            
            # Valor (si está presente, mostrarlo; si no, dejar en blanco)
            valor = respuesta_data.get('valor', '')
            contenido_respuestas += f"**Valor:** {valor}\n"

            # Observación (si está presente, mostrarla; si no, dejar en blanco)
            Observacion = respuesta_data.get('Observacion', '')
            contenido_respuestas += f"**Observación:** {Observacion}\n"

            # Si hay observación, agregarla a la sección de observaciones
            if Observacion:
                contenido_observaciones += f"{observaciones_contador}. {Observacion}\n"
                observaciones_contador += 1  # Incrementar el contador para la próxima observación                

            # Agregar un salto de línea extra después de cada subgrupo
            contenido_respuestas += "\n"

    # Reemplazar el marcador de respuestas en el documento
    for p in doc.paragraphs:
        if "RESPUESTAS_AUDITORIA" in p.text:  # Asegúrate de tener un marcador adecuado
            p.text = p.text.replace("RESPUESTAS_AUDITORIA", contenido_respuestas)
            
            # Limpiar el texto de este párrafo antes de agregarlo
            p.clear()
            
            # Dividir el contenido en partes usando los asteriscos para las negritas
            partes = contenido_respuestas.split("**")
            for i, parte in enumerate(partes):
                # Si la parte es una palabra que debe ir en negrita (índice impar)
                if i % 2 != 0:
                    # Agregar la parte en negrita
                    run = p.add_run(parte)
                    run.bold = True
                else:
                    # Si es texto normal (índice par), añadirlo sin formato
                    p.add_run(parte)

    # Reemplazar el marcador de observaciones en el documento
    for p in doc.paragraphs:
        if "OBSERVACIONES_SEGUIMIENTO" in p.text:  # Asegúrate de tener un marcador adecuado
            p.text = p.text.replace("OBSERVACIONES_SEGUIMIENTO", contenido_observaciones)
                        
    # Guardar el documento modificado
    doc.save(file_path_word)
    
    # Botón de descarga del archivo Word
    with open(file_path_word, 'rb') as word_file:
        st.download_button(label="Descargar archivo Word", data=word_file, file_name=filename_word, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.session_state["consecutivo"] += 1


# Función para reiniciar el formulario
def reiniciar_formulario():
    st.session_state["Datos Generales"] = {}
    st.session_state["responses"] = {}
    st.session_state["form"] = {}
    st.session_state["Operador"] = "SELECCIONAR"
    st.session_state.clear()

# Función para cargar la base de datos desde un archivo Excel
@st.cache_data
def load_data():
    url = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/DIRECTORIO_Operadores.txt"
    
    response = requests.get(url)
    
    if response.status_code == 200:
        return response.content  # Retorna el contenido en binario
    
    else:
        st.error(f"❌ Error al descargar el archivo: {response.status_code}")
        return None

file_content = load_data()

if file_content is not None:
    try:
        excel_data = BytesIO(file_content)  # Convierte los bytes en un archivo en memoria
        df = pd.read_csv("DIRECTORIO_Operadores.txt", sep="|", encoding ="utf-8", on_bad_lines="skip")  # Intenta leerlo con pandas
                
    except Exception as e:
        st.error(f"⚠ Error al leer el archivo con pandas: {e}")
else:
    st.warning("No se pudo descargar el archivo.")
    
# Encabezado
st.title("")

# Barra lateral: Datos generales
st.sidebar.header("Datos Generales")
if "form" not in st.session_state:st.session_state["form"] = {}
# Seleccionar operador
Operador = st.sidebar.selectbox("Operador", Operadores.keys())

# Obtener el NIT del operador seleccionado correctamente
nit_operador = Operadores.get(Operador, "DESCONOCIDO")
st.session_state["form"]["Nit_operador"] = nit_operador
st.sidebar.write(f"NIT del Operador: {nit_operador}")

# Filtrar farmacias por NIT del operador
if nit_operador.isdigit():  # Verifica si el NIT es un número válido
    farmacias_filtradas = df[df["Nit"]== int(nit_operador)]
else:
    farmacias_filtradas = pd.DataFrame()

if not farmacias_filtradas.empty:
    # Lista desplegable con nombres de farmacias
    farmacia_seleccionada = st.sidebar.selectbox("Seleccione la farmacia:", farmacias_filtradas["NOMBRE DE LA FARMACIA"].unique())

    # Obtener datos de la farmacia seleccionada
    datos_farmacia = farmacias_filtradas[farmacias_filtradas["NOMBRE DE LA FARMACIA"] == farmacia_seleccionada].iloc[0]

    # Campos prellenados pero editables
    Nit_sucursal = st.sidebar.text_input("Nit_sucursal:", datos_farmacia["COD. SUC"])
    ciudad = st.sidebar.text_input("Ciudad:", datos_farmacia["CIUDAD / MUNICIPIO"])
    direccion = st.sidebar.text_input("Dirección:", datos_farmacia["DIRECCIÓN DE CONTACTO"])
    telefono = st.sidebar.text_input("Teléfono:", datos_farmacia["TELÉFONO DE CONTACTO"])
else:
    st.sidebar.warning("No se encontraron farmacias para este operador.")

# Otros campos editables
st.session_state["form"]["Nivel y Tipo de servicio farmacéutico"] = st.sidebar.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
st.session_state["form"]["Representante legal"] = st.sidebar.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
st.session_state["form"]["Director técnico"] = st.sidebar.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
st.session_state["form"]["Auditor 1"] = st.sidebar.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
st.session_state["form"]["Auditor 2"] = st.sidebar.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
st.session_state["form"]["Fecha de Auditoria"] = st.sidebar.date_input("Fecha de Auditoria", value=pd.to_datetime(st.session_state["form"].get("Fecha de Auditoria", "2025-01-01")))
st.sidebar.radio(
    "Tipo de Droguería", ["Aliada", "Propia", "Otros"],
    key="tipo_drogueria",
    on_change=lambda: st.session_state["form"].update({"Tipo de Droguería": st.session_state["tipo_drogueria"]})
)
 
# Navegación entre grupos de preguntas
grupos = [
    "TALENTO HUMANO", "CONDICIONES LOCATIVAS",
    "DOTACION", "PROCESOS PRIORITARIOS",
    "GESTION DE LA CALIDAD", "SISTEMAS DE INFORMACION", "GESTION DE ALIADAS", "GESTION DE RIESGOS"
]

respuestas = {}

grupo_actual = st.sidebar.radio("Navegación", grupos)

#Preguntas detalladas por grupo
preguntas_por_grupo = {
    "TALENTO HUMANO":["**Concepto** El talento humano en salud, cuenta con la autorización expedida por la autoridad competente, para ejercer la profesión u ocupación.  \n\n**1.1** **Criterio a Verificar:** Cuenta y se evidencia la idoneidad del personal que labora para los diferentes procesos que se realizan.",
                      "**1.2** **Criterio a Verificar:** El servicio farmacéutico y el personal que presta el servicio en salud se encuentra registrado ante las entidades departamentales o distritales de salud con el registro especial de prestación de salud.",
                      "**Concepto** Existen contratos vigentes con el personal que labora en la institución.  \n\n**1.3** **Criterio a Verificar:**,
                      "**Concepto** Se evidencian un programa o cronograma de educación continua para el talento humano.  \n\n**1.4** **Criterio a Verificar:** Existen registro de capacitación continua y de actualización para el personal.",
                      "**Concepto** Si se  aceptan practicantes de educación superior realizados relacionados con el servicio farmacéutico (Ejm: regencia de farmacia, auxiliares de servicios farmacéuticos y química farmacéutica) en formación, debe tener procedimiento para la supervisión del personal en entrenamiento, por personal debidamente autorizado.  \n\n**1.5** **Criterio a Verificar:** Soportes de la supervisión realizada (Listados de asistencia, certificado de asignación de supervisión por la institución educativa etc).",
                      "**Concepto** Sistema de capacitación y evaluación del personal.  \n\n**1.6** **Criterio a Verificar:** La institución cuenta con cronograma de Capacitaciones del 2025.",
                      "**1.7** **Criterio a Verificar:** La institución cuenta con evidencia de la ejecución de capacitaciones durante los periodos solicitados.",
                      "**1.8** **Criterio a Verificar:** La institución cuenta con soportes de evaluación de desempeño del personal que labora en el punto evaluado.",
                      "**1.9** **Criterio a Verificar:** La institución cuenta con mecanismos de refuerzo o recapacitación para los funcionarios con desviaciones en la calidad esperada vs la observada."
                      ],
    "CONDICIONES LOCATIVAS": ["**Concepto** Los pisos, paredes y techos de todos los servicios deberán ser de fácil limpieza y estar en buenas condiciones de presentación y mantenimiento.  \n\n**2.1** **Criterio a Verificar:** Valide el documento para el proceso de limpieza: Pisos y paredes de materiales impermeables de fácil limpieza, resistentes a factores ambientales y sistema de drenaje para su fácil limpieza y sanitización.",
                              "**2.2** **Criterio a Verificar:** Valide el documento para el proceso de limpieza: Techos y cielos rasos resistentes, uniformes de fácil limpieza y desinfección.",
                              "**Concepto** Posee un sistema de ventilación natural y/o artificial que garantice la conservación adecuada de los medicamentos y dispositivos médicos.  \n\n**2.3** **Criterio a Verificar:** Valide el documento para el proceso de: Ventilación natural para los puntos.",
                              "**Concepto** Cuenta con sistema de iluminación natural o artificial para la adecuada realización de las actividades.  \n\n**2.4** **Criterio a Verificar:** Validación en punto.",
                              "**Concepto** Cuenta con tomas, interruptores y cableado protegido.  \n\n**2.5** **Criterio a Verificar:** Validación en punto: Instalaciones eléctricas presentan tomas, interruptores y cableado protegidos, también se encuentran las redes eléctricas y de datos en buen estados.",
                              "**Concepto** El establecimiento tiene identificadas y señalizadas las rutas de evacuación, las salidas de emergencia y en términos generales cada una de sus áreas en caso de accidentes.  \n\n**2.6** **Criterio a Verificar:** Validación en punto.",
                              "**Concepto** Cuenta con las áreas requeridas para la realización de los procesos del establecimiento. Diferentes para cada proceso, diferenciadas y debidamente señalizadas.  \n\n**2.7** **Criterio a Verificar:** Valide en el documento del proceso la existencia del área administrativa debidamente delimitada o señalizada.",
                              "**2.8** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área para la recepción de productos: Farmacéuticos, dispositivos médicos y productos autorizados, debidamente delimitada o señalizada.",
                              "**2.9** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área de cuarentena de medicamentos debidamente señalizada, en ella también se podrán almacenar de manera transitoria los productos retirados del mercado.",
                              "**2.10** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área del almacenamiento de medicamentos y dispositivos médicos independiente, diferenciada y señalizada.",
                              "**2.11** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área independiente y segura para el almacenamiento de medicamentos de control especial.",
                              "**2.12** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área especial, debidamente identificada, para el almacenamiento Transitorio de los medicamentos vencidos o deteriorados que deban ser técnicamente destruidos o desnaturalizados.",
                              "**2.13** **Criterio a Verificar:** Área de dispensación de medicamentos y entrega de dispositivos médicos y productos autorizados, debidamente delimitada o señalizada.",
                              "**2.14** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área para el manejo y disposición de residuos.",
                              "**2.15** **Criterio a Verificar:** Valide en el documento del proceso la existencia del Área destinada para el almacenamiento de productos rechazados, devueltos y retirados del mercado.",
                              "**Concepto** Las instalaciones se mantienen en condiciones adecuadas de conservación higiene y limpieza.  \n\n**2.16** **Criterio a Verificar:** Valide que las instalaciones esten limpias.",
                              "**Concepto** Área de aseo, donde se encuentre los implementos para la adecuada higienización.  \n\n**2.17** **Criterio a Verificar:** Recipientes limpios, suficientes, ubicados e identificados para el manejo y disposición de residuos sólidos (basura).",
                              "**2.18** **Criterio a Verificar:** Los desechos son clasificados en su fuente de generación según el código de colores: negro, blanco y verde (puntos con cafeteria) guardian o caneca roja.",
                              "**Concepto** Manejo y disposición de residuos.  \n\n**2.19** **Criterio a Verificar:** El establecimiento cuenta con el plan de gestión integral de residuos sólidos (ordinarios).",
                              "**2.20** **Criterio a Verificar:** El manejo de los residuos líquidos dentro y fuera del establecimiento, no genera molestias sanitarias, inundaciones o malos olores.",
                              "**2.21** **Criterio a Verificar:** Se cuenta con recipientes suficientes, con tapa y bolsa de colores para clasificaciòn de residuos, bien ubicados e identificados para recolección interna de los desechos según corresponda.",
                              "**2.22** **Criterio a Verificar:** Cuenta con área exclusiva  para manejo y disposición de residuos, de acuerdo con la reglamentación, protegido y en perfecto estado de mantenimiento (se lava y desinfecta cada vez que se desocupa).",
                              "**2.23** **Criterio a Verificar:** Cuentan con un contrato o certificado de recolección de residuos con empresa responsable de la recolección de residuos."
                              ],
    "DOTACION": ["**Concepto** Cuenta con la dotación y muebles exclusivos y necesarios para la adquisición, recepción, almacenamiento, conservación (como manejo de cadena de frío, medicamentos fotosensibles, higroscópicos entre otros) y dispensación de los medicamentos y dispositivos médicos para la realización de los procesos que ofrezcan, de acuerdo con las recomendaciones dadas por los fabricantes.  \n\n**3.1** **Criterio a Verificar:** ",
                 "**Concepto** Se evidencia equipos e implementos de seguridad en funcionamiento y ubicados donde tengan fácil acceso (extintores, barandas, estibas, extractoras).  \n\n**3.2** **Criterio a Verificar:** Cuentan con extintores en funcionamiento, ubicados de acuerdo a las normas de seguridad, con fecha vigente de recarga y señalizada ubicación.",
                 "**Concepto** Condiciones de temperatura y humedad, registro de condiciones ambientales, termohigrometros calibrados.  \n\n**3.3** **Criterio a Verificar:** Cuenta con equipos para la medición de condiciones de temperatura y humedad.",
                 "**3.4** **Criterio a Verificar:** Se encuentran debidamente calibrados (Certificado de calibración y hoja de vida).",
                 "**3.5** **Criterio a Verificar:** Se evidencia registro de condiciones de temperatura y humedad relativa actualizados.",
                 "**3.6** **Criterio a Verificar:** Cuenta con procedimiento para la toma de temperatura y humedad relativa, sus límites y plan de contingencia en caso de falla.",
                 "**3.7** **Criterio a Verificar:** Los equipos que se utilizan cuentan con las condiciones técnicas de calidad, respaldo y soporte técnico-científico.",
                 "**3.8** **Criterio a Verificar:** Verificar el Cronograma de mantenimiento de equipos (Neveras, Aire acondicionado, Termohigrómetros).",
                 "**Concepto** Cuenta con los equipos necesarios para cumplir con el plan de contingencia, en caso de falla eléctrica para mantener la cadena de frío.  \n\n**3.9** **Criterio a Verificar:** "
                              ],
    "PROCESOS PRIORITARIOS": ["**Concepto** Selección y compras.  \n\n**4.1** **Criterio a Verificar:** Cuentan con procedimiento documentado de selección de productos, proveedores.",
                              "**4.2** **Criterio a Verificar:** Cuentan con listado de proveedores, además de contar con los certificados de distribuir autorizado.",
                              "**4.3** **Criterio a Verificar:** En la selección de proveedores tienen en cuenta las buenas prácticas de manufactura.",
                              "**4.4** **Criterio a Verificar:** En la selección de proveedores tienen en cuenta que los productos están respaldados con registro INVIMA, cuenten con las respectivas fichas técnicas y de calidad, definida política de devolución por baja rotación y sistema de codificación de productos.",
                              "**Concepto** Adquisición y Recepción.  \n\n**4.5** **Criterio a Verificar:** Cuentan con procedimiento documentado de adquisición y recepción de medicamentos y dispositivos médicos, en esté último donde se describa la metodología de muestreo implementada.",
                              "**4.6** **Criterio a Verificar:** Se realiza recepción administrativa a todos los productos adquiridos.",
                              "**4.7** **Criterio a Verificar:** Se documenta o registran las condiciones en que se reciben los medicamentos y dispositivos médicos (formato de recepción técnica).",
                              "**4.8** **Criterio a Verificar:** Verificar las condiciones en que fueron transportados los medicamentos y dispositivos médicos antes de su recepción.",
                              "**4.9** **Criterio a Verificar:** Verificar el procedimiento de transporte de medicamentos y dispositivos médicos de CEDI a droguerías, verificando mantenimiento de la cadena de frío y condiciones de embalaje.",
                              "**4.10** **Criterio a Verificar:** El manual describe el procedimiento de devoluciones de medicamentos y dispositivos médicos.",
                              "**Concepto** Almacenamiento.   \n\n**4.11** **Criterio a Verificar:** Cuenta con proceso documentado de almacenamiento de medicamentos y dispositivos médicos.",
                              "**4.12** **Criterio a Verificar:** El sitio de almacenamiento de los medicamentos y dispositivos médicos cuenta con mecanismos que garanticen las condiciones de temperatura y humedad relativa recomendadas por el fabricante y los registros permanentes de estas variables, utilizando para ello termohigrómetros u otros instrumentos que cumplan con dichas funciones.",
                              "**4.13** **Criterio a Verificar:** Los medicamentos y dispositivos médicos se almacenan de acuerdo a un método de clasificación, que garantiza el orden, minimiza los eventos de confusión, pérdida y vencimiento durante su almacenamiento.",
                              "**4.14** **Criterio a Verificar:** Los medicamentos y dispositivos médicos están ubicados en estibas o estanterías de material sanitario, impermeable y fácil de limpiar. (estibas plásticas, no de madera).",
                              "**4.15** **Criterio a Verificar:** Cuenta con protocolos de limpieza en donde se incluyan los cronogramas para la limpieza de estanterías modulares, pero además refrigeradores- y que esté se esté implementando.",
                              "**4.16** **Criterio a Verificar:** Se evidencia un sistema de control de fechas de vencimiento de medicamentos y dispositivos médicos y una rotación adecuada de los productos farmacéuticos (primeros en vencer, primeros en salir).",
                              "**4.17** **Criterio a Verificar:** Se tienen en cuenta las condiciones dadas por el fabricante para el almacenamiento de medicamentos y dispositivos médicos.",
                              "**4.18** **Criterio a Verificar:** Se describe el procedimiento de conservación de cadena de frío que incluye el plan de contingencia.",
                              "**Concepto** Medicamentos de control especial.  \n\n**4.19** **Criterio a Verificar:** Valide documentación relacionada con: área para el almacenamiento de los medicamentos de control especial franja violeta, monopolio del estado,  es independiente, diferenciada, señalizada y se encuentra bajo llave.",
                              "**4.20** **Criterio a Verificar:** Cuenta con resolución del Fondo Nacional de Estupefacientes, para la tenencia y manejo de medicamentos de Control Especial. (Resol. 1478 de 2006).",
                              "**4.21** **Criterio a Verificar:** Cuenta con libro o kardex exclusivo para el manejo de medicamentos de control especial.",
                              "**4.22** **Criterio a Verificar:** Cumple con la presentación de informe al FNE dentro de las fechas establecidas. Validar para CEDI.",
                              "**Concepto** Entrega de medicamentos y productos pendientes.  \n\n**4.24** **Criterio a Verificar:** Se han establecido mecanismos de control para los medicamentos pendientes, establecidos en un procedimiento.",
                              "**4.25** **Criterio a Verificar:** Entregan los medicamentos  pendientes, antes de 48 horas como está establecido en la resolución 1604 de 2013.",
                              "**Concepto** Farmacovigilancia y Tecnovigilancia.  \n\n**4.26** **Criterio a Verificar:** La institución cuenta con programa de farmacovigilancia, tecnovigilancia y errores de medicación.",
                              "**4.27** **Criterio a Verificar:** Cuenta con manuales o procedimientos para la implementación de estos.",
                              "**4.28** **Criterio a Verificar:** Existe evidencia de inscripción a la red nacional de farmacovigilancia y tecnovigilancia.",
                              "**4.29** **Criterio a Verificar:** Se evidencia la presentación de informes a los entes de control."
                              ],
    "GESTION DE LA CALIDAD": ["**Concepto** Todo servicio farmacéutico, establecimiento farmacéutico o persona autorizada, tendrá la responsabilidad de desarrollar, implementar, mantener, revisar y perfeccionar un Sistema de Gestión de la Calidad Institucional, de conformidad con las leyes y demás normas sobre la materia. (Art. 13. Dec. 2200 de 2005  y art. 17 resol 1403 de 2007)  \n\n**5.1** **Criterio a Verificar:** ",
                              "**Concepto** Auditoría interna o de primer orden: El punto cuenta evidencia documental de la ejecución del "Procedimiento de auditoría interna y seguimiento a establecimientos farmacéuticos" .  \n\n**5.2** **Criterio a Verificar:** 1.1.1 Existencia de Cronograma de AuditorÍa 2025 tanto de droguerías propias.",
                              "**5.3** **Criterio a Verificar:** 1.1.2 Evidencia de la ejecución de Auditorías realizadas durante el 2024 y lo corrido del 2025.",
                              "**5.4** **Criterio a Verificar:** 1.1.3 Evidencia de la generación de plan de acción derivado de los hallazgos de las auditorías.",
                              "**5.5** **Criterio a Verificar:** 1.1.4 Evidencia de la ejecución y cumplimiento de los planes de acción establecidos por la institución derivados de las auditorías internas.",
                              "**Concepto** Auditoría Externa o de tercer orden.  \n\n**5.6** **Criterio a Verificar:** 1.2.1 La institución cuenta con el acta de la última visita realizada por los entes de Inspección vigilancia y control. (Secretaria, INVIMA, Superintendencia Nacional de Salud) año 2024 o en su defecto la última relacionada.",
                              "**5.7** **Criterio a Verificar:** 1.2.2 La institución cuenta con evidencia de la generación de Plan de acción derivado de los hallazgos de los entes inspección vigilancia y control año 2024 o en su defecto la última relacionada.",
                              "**5.8** **Criterio a Verificar:** 1.2.3 Evidencia de la ejecución y cumplimiento de los planes de acción establecidos por la institución derivados de las visitas de los entes de inspección vigilancia y control año 2024 o en su defecto la última relacionada.",
                              "**Concepto** Evaluación y seguimiento a la dispensación.  \n\n**5.9** **Criterio a Verificar:** 1.5.1. La institución cuenta con actividades de seguimiento a la adherencia del procedimiento de dispensación de medicamentos a los auxiliares de farmacia. (Método observacional).",
                              "**Concepto** Criterios y métodos necesarios para asegurar que estos procesos sean eficaces tanto en su operación como en su control.  \n\n**5.10** **Criterio a Verificar:** Se realiza seguimiento y medición a indicadores en el establecimiento o institución.",
                              "**Concepto** Puntos de control sobre los riesgos de mayor probabilidad de ocurrencia o que generen un impacto considerable en la satisfacción de las necesidades y expectativas de calidad de los usuarios, beneficiarios o destinatarios, con la participación de las personas y/o responsables de cada una de las actividades y/o procesos del servicio.  \n\n**5.11** **Criterio a Verificar:** Encuestas de satisfacción.",
                              "**Concepto** Acciones necesarias para alcanzar los resultados planificados y la mejora continua de estos procesos.  \n\n**5.12** **Criterio a Verificar:** Planes de mejora, auditorías internas.",
                              "**Concepto** Existen registros con la información de todos los medicamentos para uso humano requeridos para la prestación de los servicios que ofrece; dichos registros deben incluir el principio activo, forma farmacéutica, concentración, lote, fecha de vencimiento, presentación comercial, unidad de medida y registro sanitario vigente expedido por el INVIMA.  \n\n**5.13** **Criterio a Verificar:** "
                              ],
    "SISTEMAS DE INFORMACION": ["**Concepto** Cuenta con medios, preferiblemente computarizados que permitan el registro de adecuado de fórmulas prescritas, medicamentos de control especial y demás productos que lo requieran.  \n\n**6.1** **Criterio a Verificar:** Verificar el sistema que manejan, su correcto funcionamiento y utilización."
                                ],
    "GESTION DE ALIADAS": ["**Concepto** Aliadas.  \n\n**7.1** **Criterio a Verificar:** Durante la selección de droguerías aliadas se verifican la existencia y cumplimiento de los siguientes procedimientos: Selección, adquisición, recepción, almacenamiento, distribución y dispensación de medicamentos e insumos. Así como manejo de medicamentos de control especial y de cadena de frío.",
                           "**7.2** **Criterio a Verificar:** Se realiza verificación de la idoneidad, registro ante las entidades departamentales o distritales de salud del personal que labora en las droguerías aliadas, así como procesos de capacitación inicial y continuada.",
                           "**7.3** **Criterio a Verificar:** Se realizan auditorías periódicas a droguerías aliadas según procedimiento definido para verificar cumplimiento normativo en relación a manejo de medicamentos, insumos y dispositivos médicos."
                           ],
    "GESTION DE RIESGOS": ["**Concepto** Cuenta con matriz de riesgos la cual es de manejo y conocimiento por parte de los funcionarios.  \n\n**8.1** **Criterio a Verificar:** Solicitar la matriz de riesgos y validar que los funcionarios conozcan de su existencia e implementación."
                           ]
    }

# Opciones válidas para las respuestas
opciones_validas = ["Cumple totalmente", "Cumple parcialmente", "Incumple totalmente", "No aplica"]

# Mostrar preguntas del grupo actual
st.header(grupo_actual)
preguntas = preguntas_por_grupo.get(grupo_actual,[])

# Crear estructura de almacenamiento si no existe
if grupo_actual not in st.session_state["responses"]:
    st.session_state["responses"][grupo_actual] = {}

for pregunta in preguntas:
    if pregunta not in st.session_state["responses"][grupo_actual]:
        st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": None,
            "valor": None,
            "Observacion": "",
        }
        
       
        
    #respuestas seleccionadas previamente
    respuesta_actual = st.session_state["responses"][grupo_actual][pregunta].get("respuesta","Seleccione una opcion")
    observacion_actual = st.session_state["responses"][grupo_actual][pregunta].get("Observacion","")

    #si no hay respuesta mostrar una opcion vacia para obligar la seleccion
    Opciones = ["Selecciona una opción","Cumple totalmente", "Cumple parcialmente", "Incumple totalmente", "No Aplica"]

    # Mostrar la pregunta y permitir selección
    respuesta = st.radio(
        pregunta,
        Opciones,
        key=f"respuesta_{grupo_actual}_{pregunta}",
        index = Opciones.index(respuesta_actual) if respuesta_actual in Opciones else 0
        )
    
    #Campo de observacion
    Observacion = st.text_area(f"Observacion",
                               value=observacion_actual,
                               key=f"Observacion_{grupo_actual}_{pregunta}",
                               on_change=lambda p=pregunta: actualizar_observacion(grupo_actual,p)
                               )
    def actualizar_observacion(grupo,pregunta):
        """Actualizar la observacion en session_state al escribir en el campo de texto."""
        clave = f"Observacion_{grupo}_{pregunta}"
        st.session_state["responses"][grupo][pregunta]["Observacion"] = st.session_state[clave]
        
    # Guardar respuesta y valor
    if respuesta != "Selecciona una opción":
            valor = 100 if respuesta == "Cumple totalmente" else 50 if respuesta == "Cumple parcialmente" else 0
            if grupo_actual not in st.session_state["responses"]:
                st.session_state["responses"][grupo_actual] = {}
            st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": respuesta,
            "valor": valor,
            "Observacion": st.session_state[f"Observacion_{grupo_actual}_{pregunta}"]  # Usar el estado guardado
    }
   
# Función de validación para verificar si todas las respuestas son válidas
def validar_respuestas():
    for pregunta, data in st.session_state["responses"][grupo_actual].items():
        respuesta = data["respuesta"]
        if respuesta == "Selecciona una opción":
            st.error(f"Por favor, responde la pregunta: {pregunta}")
            return False
        if respuesta not in opciones_validas:
            st.error(f"La respuesta seleccionada para la pregunta '{pregunta}' no es válida.")
            return False
    return True

# Botón para pasar al siguiente grupo o finalizar
if st.button("Siguiente"):
    if validar_respuestas():
        st.success("Todas las preguntas fueron respondidas correctamente.")
        # Aquí podrías permitir la navegación hacia otro grupo
    else:
        st.error("Asegúrate de responder todas las preguntas correctamente.")       


# Calcular promedio
total_valores = sum(
    r["valor"] if r ["valor"] is not None else 0 for g in st.session_state["responses"].values() for r in g.values()
)
# Calcular el número total de preguntas respondidas (es decir, preguntas que tienen valor asignado)
num_preguntas_respondidas = sum(
    1 for g in st.session_state["responses"].values() 
    for r in g.values() if r["valor"] is not None
)

# Calcular el número total de preguntas posibles (en todas las hojas, sin importar si se han respondido o no)
total_preguntas_posibles = sum(len(g) for g in st.session_state["responses"].values())

# Asegurarse de que el promedio sea sobre 100, calculando sobre todas las preguntas posibles
promedio = (total_valores / (total_preguntas_posibles * 100)) * 100 if total_preguntas_posibles > 0 else 0

# Actualizamos el cálculo del promedio en la sesión para que sea persistente
st.session_state.promedio_total = round(promedio, 2)

# Mostrar el resultado
st.sidebar.metric("Promedio Total", round(promedio, 2))

# Botones
st.sidebar.button("Finalizar y Enviar", on_click=finalizar_formulario)
st.sidebar.button("Nuevo formulario", on_click=reiniciar_formulario)
st.sidebar.button("Guardar", on_click=guardar_estado)

# Entrada de usuario
consecutivo_input = st.sidebar.text_input("Ingrese el código del formulario a cargar:")

# Botón para cargar el formulario
if st.sidebar.button("Cargar formulario"):
    if consecutivo_input:
        # Indicar que se debe cargar el formulario
        st.session_state["load_form"] = True
        st.session_state["input_consecutivo"] = consecutivo_input
    else:
        st.error("Por favor, ingrese un código de formulario válido.")

# Cargar el formulario si se indicó
if st.session_state.get("load_form"):
    consecutivo_input = st.session_state.get("input_consecutivo")
    cargar_formulario_por_consecutivo(consecutivo_input)  # Llamada correcta pasando el argumento
  
