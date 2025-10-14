import streamlit as st
import os
import json
import tempfile
import io
import pandas as pd
from datetime import date
import streamlit as st
import requests
from io import BytesIO
import sys
from docx import Document

# Estilo para tonos azules y logo
st.markdown(
    """
    <style>
    /* Cambiar el color de los radio buttons */
    div[data-baseweb="radio"] label {
        color: #005A9C !important; /* Azul oscuro */
    }
    
    /* Cambiar el color del fondo cuando se selecciona */
    div[data-baseweb="radio"] input:checked + div {
        background-color: #0078D7 !important; /* Azul claro */
        color: white !important;
    }

    /* Cambiar color de fondo de la barra lateral */
    section[data-testid="stSidebar"] {
        background-color: #D9EAF7 !important; /* Azul muy claro */
    }

    /* Cambiar color de los botones */
    button {
        background-color: #0078D7 !important; /* Azul */
        color: white !important;
    }
    button:hover {
        background-color: #005A9C !important; /* Azul más oscuro */
    }
    
    /* Cambiar el color del punto del radio button sin afectar el fondo del texto */
    [role="radiogroup"] > label > div:first-child {
    background-color: #005A9C !important;  /* Azul oscuro */
    border-color: #005A9C !important;  /* Azul oscuro */
    }

    /* Mantener el fondo del texto transparente */
    [role="radiogroup"] > label > div:nth-child(2) {
    background-color: transparent !important;
    }
             
    </style>
    """,
    unsafe_allow_html=True
)


# Tu código principal
st.title("INSTRUMENTO DE VERIFICACIÓN DE SERVICIO")

# Diccionario de operadores y NITs
Operadores = {
    "SELECCIONAR": "---------------",
    "GESTOR1": "AA",
    "GESTOR2": "BB",
    "GESTOR3": "CC",
    "GESTOR4": "DD",
    "GESTOR5": "EE",
}

# Inicializar variables de sesión
if "responses" not in st.session_state:
    st.session_state["responses"] = {}
if "consecutivo" not in st.session_state:
    st.session_state["consecutivo"] = 1
if "form" not in st.session_state:
    st.session_state["form"] = {}
if "load_form" not in st.session_state:
    st.session_state["load_form"]=False   
if 'Datos Generales' not in st.session_state:
    st.session_state['Datos Generales'] = {}

# Ruta de almacenamiento de formularios
folder_path = "formularios_guardados"
os.makedirs(folder_path, exist_ok=True)

def convertir_fecha(obj):
    if isinstance(obj, date):
        return obj.isoformat()  # Convierte a cadena 'YYYY-MM-DD'
    raise TypeError("Tipo no serializable")


# Función para guardar el estado actual del formulario
def guardar_estado():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = date.today().isoformat()
    consecutivo = f"{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}.json"
    file_path = os.path.join(folder_path, filename)
    

    # Crear contenido JSON
    data = {
        "form": st.session_state["form"],
        "responses": st.session_state["responses"],
        "consecutivo": st.session_state["consecutivo"],
        "load_form": st.session_state["load_form"]
    }
    
    
    # Guardar en un archivo temporal para permitir la descarga
    with open(file_path, "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=4, default=convertir_fecha)

    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")

    # Botón para descargar
    with open(file_path, "r",encoding="utf-8") as file:
        st.download_button(
            label="Descargar formulario",
            data=file.read(),
            file_name=filename,
            mime="application/json",
        )

# Función para cargar formulario desde un archivo JSON basado en un consecutivo
def cargar_formulario_por_consecutivo(consecutivo_input):
    # Ahora la función acepta un argumento
    archivo_a_cargar = f"Formulario_{str(consecutivo_input)}.json"
    file_path = os.path.join(folder_path, archivo_a_cargar)

    if os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Actualizar session_state con los datos cargados
            st.session_state["Datos Generales"] = data.get("Datos Generales",{})
            st.session_state["form"] = data.get("form", {})
            st.session_state["responses"] = data.get("responses", {})
            st.session_state["consecutivo"] = data.get("consecutivo", 1)
            st.session_state["form"]["observacion"] = st.session_state["form"].get("observacion", "")
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Variables adicionales de datos generales
            operador = st.session_state["form"].get("Operador", "SELECCIONAR")
            nit_operador = operadores.get(operador, "DESCONOCIDO")
            farmacia_seleccionada = st.session_state["form"].get("farmacia_seleccionada", "NO SELECCIONADA")
            Nit_sucursal = farmacias_filtradas.get("Nit_sucursal:", datos_farmacia.get("COD. SUC", "DESCONOCIDO"))
            ciudad = st.session_state["form"].get("ciudad", "CIUDAD NO ESPECIFICADA")
            direccion = st.session_state["form"].get("direccion", "DIRECCIÓN NO ESPECIFICADA")
            telefono = st.session_state["form"].get("telefono", "TELEFONO NO ESPECIFICADO")

            # Mostrar los campos en el formulario
            st.text_input("Operador", value=operador)
            st.text_input("Nit del Operador", value=nit_operador)
            st.text_input("Farmacia Seleccionada", value=farmacia_seleccionada)
            st.text_input("Nit de la Sucursal", value=Nit_sucursal)
            st.text_input("Ciudad", value=ciudad)
            st.text_input("Dirección", value=direccion)
            st.text_input("Teléfono", value=telefono)

            # Mostrar otros campos del formulario
            st.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
            st.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
            st.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
            st.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
            st.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
            st.text_input("Fecha de Auditoría", value=st.session_state["form"].get("Fecha de Auditoria", ""))

            # Mostrar la observación en un área de texto
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Escribir los datos generales y respuestas para la visualización
            Contenido = ""
            for grupo, preguntas in st.session_state["responses"].items():
                for pregunta, respuesta in preguntas.items():
                    Contenido += (
                        f"{operador}|{nit_operador}|"
                        f"{farmacia_seleccionada}|"
                        f"{Nit_sucursal}|"
                        f"{ciudad}|"
                        f"{direccion}|"
                        f"{telefono}|"
                        f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                        f"{st.session_state['form'].get('Representante legal', '')}|"
                        f"{st.session_state['form'].get('Director técnico', '')}|"
                        f"{st.session_state['form'].get('Auditor 1', '')}|"
                        f"{st.session_state['form'].get('Auditor 2', '')}|"
                        f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    )
            st.write("Contenido generado:")
            st.write(Contenido)

            st.success(f"Formulario {archivo_a_cargar} cargado correctamente.")
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
    else:
        st.error(f"No se encontró un archivo con el consecutivo {consecutivo_input}.")


# Función para finalizar y generar el archivo TXT
def finalizar_formulario():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    nit_operador = Operadores.get(Operador, "DESCONOCIDO")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = "2025"
    consecutivo = f"{fecha_auditoria[:4]}_{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}.txt"
    filename_word = f"Acta_Seguimiento_{farmacia_seleccionada}.docx"
    folder_path = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/Acta_Seguimiento_PUBLICO.docx"

    # Descargar el archivo .docx desde GitHub
    response = requests.get(folder_path)
    if response.status_code == 200:
        # Cargar el archivo .docx directamente desde la respuesta
        doc = Document(BytesIO(response.content))
    else:
        st.error(f"Error al descargar el archivo: {response.status_code}")
        return

    # Continuar con el resto de tu lógica
    file_path = os.path.join(folder_path, filename)
    file_path_word = os.path.join(folder_path, filename_word)

  
     # Escribir cabecera
    Contenido = "Operador|Nit operador|Nombre de la droguería o farmacia|Nit_sucursal|Ciudad|Dirección|Teléfono|"
    Contenido += "Nivel y Tipo de servicio farmacéutico|Representante legal|Director técnico|"
    Contenido += "Auditor 1|Auditor 2|Fecha de Auditoria|Grupo de Pregunta|Respuesta|Valor|Observacion|Subgrupo de pregunta\n"

    Operador = st.session_state["form"].get("Operador","")
    nit_operador = st.session_state["form"].get("Nit_operador", "DESCONOCIDO")
    nombre_farmacia = st.session_state["form"].get("Nombre")
    nivel_servicio = st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", "")
    representante_legal = st.session_state["form"].get("Representante legal", "")
    director_tecnico = st.session_state["form"].get("Director técnico", "")
    auditor_2 = st.session_state["form"].get("Auditor 2", "")
    Observacion = st.session_state["form"].get("Observacion", "")
    tipo_drogueria = st.session_state["form"].get("Tipo de Droguería", "")
    
        
     # Escribir datos generales y respuestas
    for grupo, preguntas in st.session_state["responses"].items():
            for pregunta, respuesta in preguntas.items():
                Contenido +=(
                    f"{Operador}|{nit_operador}|"
                    f"{farmacia_seleccionada}|"
                    f"{Nit_sucursal}|"
                    f"{ciudad}|"
                    f"{direccion}|"
                    f"{telefono}|"
                    f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                    f"{st.session_state['form'].get('Representante legal', '')}|"
                    f"{st.session_state['form'].get('Director técnico', '')}|"
                    f"{st.session_state['form'].get('Auditor 1', '')}|"
                    f"{st.session_state['form'].get('Auditor 2', '')}|"
                    f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    f"{grupo}|{respuesta['respuesta']}|{respuesta['valor']}|"
                    f"{Observacion}|"
                    f"{pregunta}\n"
                )
    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")    
            
    # Convertir el contenido a bytes
    Contenido_bytes = Contenido.encode("utf-8")
    
    #boton de descarga
    st.download_button(label="Descargar archivo txt",
                   data=Contenido,
                   file_name=filename,
                   mime="text/plain"
                   )

    # Crear archivo txt
    os.makedirs(folder_path, exist_ok=True)
    with open(file_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(Contenido) 

    # Rellenar los campos de la plantilla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("OPERADOR", Operador)
                if "NIT_OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_OPERADOR", nit_operador)
                if "NOMBRE_FARMACIA" in cell.text.strip():
                    cell.text = cell.text.replace("NOMBRE_FARMACIA", farmacia_seleccionada)
                if "NIT_SUCURSAL" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_SUCURSAL", Nit_sucursal)
                if "CIUDAD" in cell.text.strip():
                    cell.text = cell.text.replace("CIUDAD", ciudad)
                if "DIRECCION" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECCION", direccion)
                if "TELEFONO" in cell.text.strip():
                    cell.text = cell.text.replace("TELEFONO", telefono)
                if "NIVEL_SERVICIO" in cell.text.strip():
                    cell.text = cell.text.replace("NIVEL_SERVICIO", nivel_servicio)
                if "REPRESENTANTE_LEGAL" in cell.text.strip():
                    cell.text = cell.text.replace("REPRESENTANTE_LEGAL", representante_legal)
                if "DIRECTOR_TECNICO" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECTOR_TECNICO", director_tecnico)
                if "AUDITOR_1" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_1", Auditor)
                if "AUDITOR_2" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_2", auditor_2)
                if "FECHA_AUDITORIA" in cell.text.strip():
                    cell.text = cell.text.replace("FECHA_AUDITORIA", fecha_auditoria)
                if "TIPO_DROGERIA" in cell.text.strip():
                    cell.text = cell.text.replace("TIPO_DROGERIA", tipo_drogueria)    
                    
    # Crear un string que contendrá todo el contenido a insertar en el documento
    contenido_respuestas = ""
    contenido_observaciones = ""
    
    # Recorrer los grupos y subgrupos de preguntas
    observaciones_contador = 1
    for grupo, subgrupos in st.session_state["responses"].items():
        # Primero agregamos el nombre del grupo, que se repetirá solo una vez
        contenido_respuestas += f"**{grupo}**\n"  # Agregar el grupo en negrita
        
        for subgrupo, respuesta_data in subgrupos.items():
            # Formatear el texto para cada grupo y subgrupo
            contenido_respuestas += f"{subgrupo}\n\n"

            # Respuesta
            respuesta = respuesta_data.get('respuesta', 'sin respuesta')
            contenido_respuestas += f"**Respuesta:** {respuesta}\n"
            
            # Valor (si está presente, mostrarlo; si no, dejar en blanco)
            valor = respuesta_data.get('valor', '')
            contenido_respuestas += f"**Valor:** {valor}\n"

            # Observación (si está presente, mostrarla; si no, dejar en blanco)
            Observacion = respuesta_data.get('Observacion', '')
            contenido_respuestas += f"**Observación:** {Observacion}\n"

            # Si hay observación, agregarla a la sección de observaciones
            if Observacion:
                contenido_observaciones += f"{observaciones_contador}. {Observacion}\n"
                observaciones_contador += 1  # Incrementar el contador para la próxima observación                

            # Agregar un salto de línea extra después de cada subgrupo
            contenido_respuestas += "\n"

    # Reemplazar el marcador de respuestas en el documento
    for p in doc.paragraphs:
        if "RESPUESTAS_AUDITORIA" in p.text:  # Asegúrate de tener un marcador adecuado
            p.text = p.text.replace("RESPUESTAS_AUDITORIA", contenido_respuestas)
            
            # Limpiar el texto de este párrafo antes de agregarlo
            p.clear()
            
            # Dividir el contenido en partes usando los asteriscos para las negritas
            partes = contenido_respuestas.split("**")
            for i, parte in enumerate(partes):
                # Si la parte es una palabra que debe ir en negrita (índice impar)
                if i % 2 != 0:
                    # Agregar la parte en negrita
                    run = p.add_run(parte)
                    run.bold = True
                else:
                    # Si es texto normal (índice par), añadirlo sin formato
                    p.add_run(parte)

    # Reemplazar el marcador de observaciones en el documento
    for p in doc.paragraphs:
        if "OBSERVACIONES_SEGUIMIENTO" in p.text:  # Asegúrate de tener un marcador adecuado
            p.text = p.text.replace("OBSERVACIONES_SEGUIMIENTO", contenido_observaciones)
                        
    # Guardar el documento modificado
    doc.save(file_path_word)
    
    # Botón de descarga del archivo Word
    with open(file_path_word, 'rb') as word_file:
        st.download_button(label="Descargar archivo Word", data=word_file, file_name=filename_word, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.session_state["consecutivo"] += 1


# Función para reiniciar el formulario
def reiniciar_formulario():
    st.session_state["Datos Generales"] = {}
    st.session_state["responses"] = {}
    st.session_state["form"] = {}
    st.session_state["Operador"] = "SELECCIONAR"
    st.session_state.clear()

# Función para cargar la base de datos desde un archivo Excel
@st.cache_data
def load_data():
    url = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/DIRECTORIO_GESTORES_PUBLICO.txt"
    
    response = requests.get(url)
    
    if response.status_code == 200:
        return response.content  # Retorna el contenido en binario
    
    else:
        st.error(f"❌ Error al descargar el archivo: {response.status_code}")
        return None

file_content = load_data()

if file_content is not None:
    try:
        excel_data = BytesIO(file_content)  # Convierte los bytes en un archivo en memoria
        df = pd.read_csv("DIRECTORIO_GESTORES_PUBLICO.txt", sep="|", encoding ="utf-8", on_bad_lines="skip")  # Intenta leerlo con pandas
                
    except Exception as e:
        st.error(f"⚠ Error al leer el archivo con pandas: {e}")
else:
    st.warning("No se pudo descargar el archivo.")
    
# Encabezado
st.title("")

# Barra lateral: Datos generales
st.sidebar.header("Datos Generales")
if "form" not in st.session_state:st.session_state["form"] = {}
# Seleccionar operador
Operador = st.sidebar.selectbox("Operador", Operadores.keys())

# Obtener el NIT del operador seleccionado correctamente
nit_operador = Operadores.get(Operador, "DESCONOCIDO")
st.session_state["form"]["Nit_operador"] = nit_operador
st.sidebar.write(f"NIT del Operador: {nit_operador}")

# Filtrar farmacias por NIT del operador
if nit_operador.isdigit():  # Verifica si el NIT es un número válido
    farmacias_filtradas = df[df["Nit"]== int(nit_operador)]
else:
    farmacias_filtradas = pd.DataFrame()

if not farmacias_filtradas.empty:
    # Lista desplegable con nombres de farmacias
    farmacia_seleccionada = st.sidebar.selectbox("Seleccione la farmacia:", farmacias_filtradas["NOMBRE DE LA FARMACIA"].unique())

    # Obtener datos de la farmacia seleccionada
    datos_farmacia = farmacias_filtradas[farmacias_filtradas["NOMBRE DE LA FARMACIA"] == farmacia_seleccionada].iloc[0]

    # Campos prellenados pero editables
    Nit_sucursal = st.sidebar.text_input("Nit_sucursal:", datos_farmacia["COD. SUC"])
    ciudad = st.sidebar.text_input("Ciudad:", datos_farmacia["CIUDAD / MUNICIPIO"])
    direccion = st.sidebar.text_input("Dirección:", datos_farmacia["DIRECCIÓN DE CONTACTO"])
    telefono = st.sidebar.text_input("Teléfono:", datos_farmacia["TELÉFONO DE CONTACTO"])
else:
    st.sidebar.warning("No se encontraron farmacias para este operador.")

# Otros campos editables
st.session_state["form"]["Nivel y Tipo de servicio farmacéutico"] = st.sidebar.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
st.session_state["form"]["Representante legal"] = st.sidebar.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
st.session_state["form"]["Director técnico"] = st.sidebar.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
st.session_state["form"]["Auditor 1"] = st.sidebar.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
st.session_state["form"]["Auditor 2"] = st.sidebar.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
st.session_state["form"]["Fecha de Auditoria"] = st.sidebar.date_input("Fecha de Auditoria", value=pd.to_datetime(st.session_state["form"].get("Fecha de Auditoria", "2025-01-01")))
st.sidebar.radio(
    "Tipo de Droguería", ["Aliada", "Propia", "Otros"],
    key="tipo_drogueria",
    on_change=lambda: st.session_state["form"].update({"Tipo de Droguería": st.session_state["tipo_drogueria"]})
)
 
# Navegación entre grupos de preguntas
grupos = [
    "CONCEPTO 1", "CONCEPTO 2", "CONCEPTO 3"
]

respuestas = {}

grupo_actual = st.sidebar.radio("Navegación", grupos)

#Preguntas detalladas por grupo
preguntas_por_grupo = {
    "CONCEPTO 1":["**Concepto** PREGUNTA\n\n**1.1** **Criterio a Verificar:** DESCRIPCION",
                          "**Concepto** PREGUNTA  \n\n**1.2** **Criterio a Verificar:** DESCRIPCION",
                          "**Concepto** PREGUNTA  \n\n**1.3** **Criterio a Verificar:** DESCRIPCION"
                          ],
    "CONCEPTO 2":["**Concepto** PREGUNTA  \n\n**2.1** **Criterio a Verificar:** DESCRIPCION",
                      "**Concepto** PREGUNTA  \n\n**2.2** **Criterio a Verificar:** DESCRIPCION",
                      "**Concepto** PREGUNTA  \n\n**2.3** **Criterio a Verificar:** DESCRIPCION"
                      ],
    "CONCEPTO 3": ["**Concepto** PREGUNTA  \n\n**3.1** **Criterio a Verificar:** DESCRIPCION",
                        "**Concepto** PREGUNTA  \n\n**3.2** **Criterio a Verificar:** DESCRIPCION",
                        "**Concepto** PREGUNTA  \n\n**3.3** **Criterio a Verificar:** DESCRIPCION"
                        ],
    }

# Opciones válidas para las respuestas
opciones_validas = ["Cumple totalmente", "Cumple parcialmente", "Incumple totalmente", "No aplica"]

# Mostrar preguntas del grupo actual
st.header(grupo_actual)
preguntas = preguntas_por_grupo.get(grupo_actual,[])

# Crear estructura de almacenamiento si no existe
if grupo_actual not in st.session_state["responses"]:
    st.session_state["responses"][grupo_actual] = {}

for pregunta in preguntas:
    if pregunta not in st.session_state["responses"][grupo_actual]:
        st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": None,
            "valor": None,
            "Observacion": "",
        }
        
       
        
    #respuestas seleccionadas previamente
    respuesta_actual = st.session_state["responses"][grupo_actual][pregunta].get("respuesta","Seleccione una opcion")
    observacion_actual = st.session_state["responses"][grupo_actual][pregunta].get("Observacion","")

    #si no hay respuesta mostrar una opcion vacia para obligar la seleccion
    Opciones = ["Selecciona una opción","Cumple totalmente", "Cumple parcialmente", "Incumple totalmente", "No Aplica"]

    # Mostrar la pregunta y permitir selección
    respuesta = st.radio(
        pregunta,
        Opciones,
        key=f"respuesta_{grupo_actual}_{pregunta}",
        index = Opciones.index(respuesta_actual) if respuesta_actual in Opciones else 0
        )
    
    #Campo de observacion
    Observacion = st.text_area(f"Observacion",
                               value=observacion_actual,
                               key=f"Observacion_{grupo_actual}_{pregunta}",
                               on_change=lambda p=pregunta: actualizar_observacion(grupo_actual,p)
                               )
    def actualizar_observacion(grupo,pregunta):
        """Actualizar la observacion en session_state al escribir en el campo de texto."""
        clave = f"Observacion_{grupo}_{pregunta}"
        st.session_state["responses"][grupo][pregunta]["Observacion"] = st.session_state[clave]
        
    # Guardar respuesta y valor
    if respuesta != "Selecciona una opción":
            valor = 100 if respuesta == "Cumple totalmente" else 50 if respuesta == "Cumple parcialmente" else 0
            if grupo_actual not in st.session_state["responses"]:
                st.session_state["responses"][grupo_actual] = {}
            st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": respuesta,
            "valor": valor,
            "Observacion": st.session_state[f"Observacion_{grupo_actual}_{pregunta}"]  # Usar el estado guardado
    }
   
# Función de validación para verificar si todas las respuestas son válidas
def validar_respuestas():
    for pregunta, data in st.session_state["responses"][grupo_actual].items():
        respuesta = data["respuesta"]
        if respuesta == "Selecciona una opción":
            st.error(f"Por favor, responde la pregunta: {pregunta}")
            return False
        if respuesta not in opciones_validas:
            st.error(f"La respuesta seleccionada para la pregunta '{pregunta}' no es válida.")
            return False
    return True

# Botón para pasar al siguiente grupo o finalizar
if st.button("Siguiente"):
    if validar_respuestas():
        st.success("Todas las preguntas fueron respondidas correctamente.")
        # Aquí podrías permitir la navegación hacia otro grupo
    else:
        st.error("Asegúrate de responder todas las preguntas correctamente.")       


# Calcular promedio
total_valores = sum(
    r["valor"] if r ["valor"] is not None else 0 for g in st.session_state["responses"].values() for r in g.values()
)
# Calcular el número total de preguntas respondidas (es decir, preguntas que tienen valor asignado)
num_preguntas_respondidas = sum(
    1 for g in st.session_state["responses"].values() 
    for r in g.values() if r["valor"] is not None
)

# Calcular el número total de preguntas posibles (en todas las hojas, sin importar si se han respondido o no)
total_preguntas_posibles = sum(len(g) for g in st.session_state["responses"].values())

# Asegurarse de que el promedio sea sobre 100, calculando sobre todas las preguntas posibles
promedio = (total_valores / (total_preguntas_posibles * 100)) * 100 if total_preguntas_posibles > 0 else 0

# Actualizamos el cálculo del promedio en la sesión para que sea persistente
st.session_state.promedio_total = round(promedio, 2)

# Mostrar el resultado
st.sidebar.metric("Promedio Total", round(promedio, 2))

# Botones
st.sidebar.button("Finalizar y Enviar", on_click=finalizar_formulario)
st.sidebar.button("Nuevo formulario", on_click=reiniciar_formulario)
st.sidebar.button("Guardar", on_click=guardar_estado)

# Entrada de usuario
consecutivo_input = st.sidebar.text_input("Ingrese el código del formulario a cargar:")

# Botón para cargar el formulario
if st.sidebar.button("Cargar formulario"):
    if consecutivo_input:
        # Indicar que se debe cargar el formulario
        st.session_state["load_form"] = True
        st.session_state["input_consecutivo"] = consecutivo_input
    else:
        st.error("Por favor, ingrese un código de formulario válido.")

# Cargar el formulario si se indicó
if st.session_state.get("load_form"):
    consecutivo_input = st.session_state.get("input_consecutivo")
    cargar_formulario_por_consecutivo(consecutivo_input)  # Llamada correcta pasando el argumento
    st.session_state["load_form"] = False  # Restablecer el indicador
