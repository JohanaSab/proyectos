import streamlit as st
import os
import json
import tempfile
import io
import pandas as pd
from datetime import date
import streamlit as st
import requests
from io import BytesIO
import sys
from docx import Document

# Estilo para tonos azules y logo
st.markdown(
    """
    <style>
    /* Cambiar el color de los radio buttons */
    div[data-baseweb="radio"] label {
        color: #005A9C !important; /* Azul oscuro */
    }
    
    /* Cambiar el color del fondo cuando se selecciona */
    div[data-baseweb="radio"] input:checked + div {
        background-color: #0078D7 !important; /* Azul claro */
        color: white !important;
    }

    /* Cambiar color de fondo de la barra lateral */
    section[data-testid="stSidebar"] {
        background-color: #D9EAF7 !important; /* Azul muy claro */
    }

    /* Cambiar color de los botones */
    button {
        background-color: #0078D7 !important; /* Azul */
        color: white !important;
    }
    button:hover {
        background-color: #005A9C !important; /* Azul más oscuro */
    }
    
    /* Cambiar el color del punto del radio button sin afectar el fondo del texto */
    [role="radiogroup"] > label > div:first-child {
    background-color: #005A9C !important;  /* Azul oscuro */
    border-color: #005A9C !important;  /* Azul oscuro */
    }

    /* Mantener el fondo del texto transparente */
    [role="radiogroup"] > label > div:nth-child(2) {
    background-color: transparent !important;
    }
             
    </style>
    """,
    unsafe_allow_html=True
)

# Agregar el logo
logo_path = "logo2.png"  # Cambia la ruta al archivo de tu logo
st.sidebar.image(logo_path, use_container_width=True)

# Tu código principal
st.title("INSTRUMENTO DE VERIFICACIÓN DE SERVICIO FARMACÉUTICO")

# Diccionario de operadores y NITs
Operadores = {
    "SELECCIONAR": "---------------",
    "COHAN": "890985122",
    "GENHOSPI": "900331412",
    "INSUMEDIC": "900716566",
    "MEDICINA Y TECNOLOGIA EN SALUD MYT": "900057926",
    "SUMINISTROS Y DOTACIONES": "802000608",
    "DISFARMA": "900580962",
    "CRUZ VERDE": "800149695",
    "DISCOLMETS": "828002423",
    "HUMSALUD": "900994906",
    "ETICOS": "892300678",
    "JEZA": "901048992",
    "MEDISFARMA": "901148499",
    "DISTRIMEVD": "901433283",
    "CYA": "860029216",
    "AUDIFARMA": "816001182",
    "FARMACIA INST": "900285194",
    "INHAPRO": "901200716",
    "PHARMASAN": "900249425",
    "CLINICA EL CARMEN": "800149384",
    "VALENTECH PHARMA COLOMBIA SAS": "900677118",
    "LIGA CONTRA EL CANCER SECCIONAL RISARALDA": "891408586",
    "CENTRO ONCOLOGICO DE ANTIOQUIA SAS": "900236850",
    "UNIDAD HEMATO ONCOLOGICA Y DE RADIOTERAPIA DEL CARIBE": "900095504",
    "CLINICA LA ERMITA DE CARTAGENA": "900491883",
    "UNION TEMPORAL ONCOLOGICA ISNOOS SN": "901776252",
    "SANOFI AVENTIS DE COLOMBIA SA": "830010337",
    "NEXT PHARMA COLOMBIA SAS": "900382525",
    "EXELTIS SAS": "900716452",
    "FORPRESALUD": "804008792",
    "MACROMED": "830107855",
    "MEDYTEC": "900057926",
    "OFFIMEDICAS": "900098550",
    "ETTICOS": "892300678",
    "MULTICARE": "901671387"
}

# Inicializar variables de sesión
if "responses" not in st.session_state:
    st.session_state["responses"] = {}
if "consecutivo" not in st.session_state:
    st.session_state["consecutivo"] = 1
if "form" not in st.session_state:
    st.session_state["form"] = {}
if "load_form" not in st.session_state:
    st.session_state["load_form"]=False   
if 'Datos Generales' not in st.session_state:
    st.session_state['Datos Generales'] = {}

# Ruta de almacenamiento de formularios
folder_path = "formularios_guardados"
os.makedirs(folder_path, exist_ok=True)

def convertir_fecha(obj):
    if isinstance(obj, date):
        return obj.isoformat()  # Convierte a cadena 'YYYY-MM-DD'
    raise TypeError("Tipo no serializable")


# Función para guardar el estado actual del formulario
def guardar_estado():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    nit_operador = Operadores.get(Operador, "DESCONOCIDO")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = date.today().isoformat()
    consecutivo = f"{fecha_auditoria[:4]}_{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}_{fecha_auditoria}_{consecutivo}.json"
    file_path = os.path.join(folder_path, filename)
    

    # Crear contenido JSON
    data = {
        "form": st.session_state["form"],
        "responses": st.session_state["responses"],
        "consecutivo": st.session_state["consecutivo"],
        "load_form": st.session_state["load_form"]
    }
    
    
    # Guardar en un archivo temporal para permitir la descarga
    with open(file_path, "w", encoding="utf-8") as file:
        json.dump(data, file, ensure_ascii=False, indent=4, default=convertir_fecha)

    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")

    # Botón para descargar
    with open(file_path, "r",encoding="utf-8") as file:
        st.download_button(
            label="Descargar formulario",
            data=file.read(),
            file_name=filename,
            mime="application/json",
        )

# Función para cargar formulario desde un archivo JSON basado en un consecutivo
def cargar_formulario_por_consecutivo(consecutivo_input):
    # Ahora la función acepta un argumento
    archivo_a_cargar = f"Formulario_{str(consecutivo_input)}.json"
    file_path = os.path.join(folder_path, archivo_a_cargar)

    if os.path.exists(file_path):
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                data = json.load(file)

            # Actualizar session_state con los datos cargados
            st.session_state["Datos Generales"] = data.get("Datos Generales",{})
            st.session_state["form"] = data.get("form", {})
            st.session_state["responses"] = data.get("responses", {})
            st.session_state["consecutivo"] = data.get("consecutivo", 1)
            st.session_state["form"]["observacion"] = st.session_state["form"].get("observacion", "")
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Variables adicionales de datos generales
            operador = st.session_state["form"].get("Operador", "SELECCIONAR")
            nit_operador = operadores.get(operador, "DESCONOCIDO")
            farmacia_seleccionada = st.session_state["form"].get("farmacia_seleccionada", "NO SELECCIONADA")
            Nit_sucursal = farmacias_filtradas.get("Nit_sucursal:", datos_farmacia.get("COD. SUC", "DESCONOCIDO"))
            ciudad = st.session_state["form"].get("ciudad", "CIUDAD NO ESPECIFICADA")
            direccion = st.session_state["form"].get("direccion", "DIRECCIÓN NO ESPECIFICADA")
            telefono = st.session_state["form"].get("telefono", "TELEFONO NO ESPECIFICADO")

            # Mostrar los campos en el formulario
            st.text_input("Operador", value=operador)
            st.text_input("Nit del Operador", value=nit_operador)
            st.text_input("Farmacia Seleccionada", value=farmacia_seleccionada)
            st.text_input("Nit de la Sucursal", value=Nit_sucursal)
            st.text_input("Ciudad", value=ciudad)
            st.text_input("Dirección", value=direccion)
            st.text_input("Teléfono", value=telefono)

            # Mostrar otros campos del formulario
            st.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
            st.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
            st.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
            st.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
            st.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
            st.text_input("Fecha de Auditoría", value=st.session_state["form"].get("Fecha de Auditoria", ""))

            # Mostrar la observación en un área de texto
            st.text_area("Observación", value=st.session_state["form"].get("observacion", ""))

            # Escribir los datos generales y respuestas para la visualización
            Contenido = ""
            for grupo, preguntas in st.session_state["responses"].items():
                for pregunta, respuesta in preguntas.items():
                    Contenido += (
                        f"{operador}|{nit_operador}|"
                        f"{farmacia_seleccionada}|"
                        f"{Nit_sucursal}|"
                        f"{ciudad}|"
                        f"{direccion}|"
                        f"{telefono}|"
                        f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                        f"{st.session_state['form'].get('Representante legal', '')}|"
                        f"{st.session_state['form'].get('Director técnico', '')}|"
                        f"{st.session_state['form'].get('Auditor 1', '')}|"
                        f"{st.session_state['form'].get('Auditor 2', '')}|"
                        f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    )
            st.write("Contenido generado:")
            st.write(Contenido)

            st.success(f"Formulario {archivo_a_cargar} cargado correctamente.")
        except Exception as e:
            st.error(f"Error al cargar el archivo: {e}")
    else:
        st.error(f"No se encontró un archivo con el consecutivo {consecutivo_input}.")


# Función para finalizar y generar el archivo TXT
def finalizar_formulario():
    Operador = st.session_state["form"].get("Operador", "SELECCIONAR")
    nit_operador = Operadores.get(Operador, "DESCONOCIDO")
    Nit_sucursal=farmacias_filtradas.get("Nit_sucursal:", datos_farmacia["COD. SUC"])
    farmacia_seleccionada = farmacias_filtradas.get("farmacia_seleccionada", datos_farmacia["NOMBRE DE LA FARMACIA"])
    fecha_auditoria = "2025"
    consecutivo = f"{fecha_auditoria[:4]}_{st.session_state['consecutivo']}"
    Auditor = st.session_state["form"].get("Auditor 1")
    filename = f"Formulario_{farmacia_seleccionada}_{fecha_auditoria}_{consecutivo}.txt"
    filename_word = f"Acta_Seguimiento_{farmacia_seleccionada}_{fecha_auditoria}_{consecutivo}.docx"
    folder_path = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/Acta_Seguimiento.docx"

    # Descargar el archivo .docx desde GitHub
    response = requests.get(folder_path)
    if response.status_code == 200:
        # Cargar el archivo .docx directamente desde la respuesta
        doc = Document(BytesIO(response.content))
    else:
        st.error(f"Error al descargar el archivo: {response.status_code}")
        return

    # Continuar con el resto de tu lógica
    file_path = os.path.join(folder_path, filename)
    file_path_word = os.path.join(folder_path, filename_word)

  
     # Escribir cabecera
    Contenido = "Operador|Nit operador|Nombre de la droguería o farmacia|Nit_sucursal|Ciudad|Dirección|Teléfono|"
    Contenido += "Nivel y Tipo de servicio farmacéutico|Representante legal|Director técnico|"
    Contenido += "Auditor 1|Auditor 2|Fecha de Auditoria|Grupo de Pregunta|Respuesta|Valor|Observacion|Subgrupo de pregunta\n"

    Operador = st.session_state["form"].get("Operador","")
    nit_operador = st.session_state["form"].get("Nit_operador", "DESCONOCIDO")
    nombre_farmacia = st.session_state["form"].get("Nombre")
    nivel_servicio = st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", "")
    representante_legal = st.session_state["form"].get("Representante legal", "")
    director_tecnico = st.session_state["form"].get("Director técnico", "")
    auditor_2 = st.session_state["form"].get("Auditor 2", "")
    Observacion = st.session_state["form"].get("Observacion", "")
    tipo_drogueria = st.session_state["form"].get("Tipo de Droguería", "")
    
        
     # Escribir datos generales y respuestas
    for grupo, preguntas in st.session_state["responses"].items():
            for pregunta, respuesta in preguntas.items():
                Contenido +=(
                    f"{Operador}|{nit_operador}|"
                    f"{farmacia_seleccionada}|"
                    f"{Nit_sucursal}|"
                    f"{ciudad}|"
                    f"{direccion}|"
                    f"{telefono}|"
                    f"{st.session_state['form'].get('Nivel y Tipo de servicio farmacéutico', '')}|"
                    f"{st.session_state['form'].get('Representante legal', '')}|"
                    f"{st.session_state['form'].get('Director técnico', '')}|"
                    f"{st.session_state['form'].get('Auditor 1', '')}|"
                    f"{st.session_state['form'].get('Auditor 2', '')}|"
                    f"{st.session_state['form'].get('Fecha de Auditoria', '')}|"
                    f"{grupo}|{respuesta['respuesta']}|{respuesta['valor']}|"
                    f"{Observacion}|"
                    f"{pregunta}\n"
                )
    st.session_state["consecutivo"] += 1
    st.success(f"Formulario guardado como: {filename}")    
            
    # Convertir el contenido a bytes
    Contenido_bytes = Contenido.encode("utf-8")
    
    #boton de descarga
    st.download_button(label="Descargar archivo txt",
                   data=Contenido,
                   file_name=filename,
                   mime="text/plain"
                   )

    # Crear archivo txt
    os.makedirs(folder_path, exist_ok=True)
    with open(file_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(Contenido) 

    # Rellenar los campos de la plantilla
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("OPERADOR", nit_operador)
                if "NIT_OPERADOR" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_OPERADOR", nit_operador)
                if "NOMBRE_FARMACIA" in cell.text.strip():
                    cell.text = cell.text.replace("NOMBRE_FARMACIA", farmacia_seleccionada)
                if "NIT_SUCURSAL" in cell.text.strip():
                    cell.text = cell.text.replace("NIT_SUCURSAL", Nit_sucursal)
                if "CIUDAD" in cell.text.strip():
                    cell.text = cell.text.replace("CIUDAD", ciudad)
                if "DIRECCION" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECCION", direccion)
                if "TELEFONO" in cell.text.strip():
                    cell.text = cell.text.replace("TELEFONO", telefono)
                if "NIVEL_SERVICIO" in cell.text.strip():
                    cell.text = cell.text.replace("NIVEL_SERVICIO", nivel_servicio)
                if "REPRESENTANTE_LEGAL" in cell.text.strip():
                    cell.text = cell.text.replace("REPRESENTANTE_LEGAL", representante_legal)
                if "DIRECTOR_TECNICO" in cell.text.strip():
                    cell.text = cell.text.replace("DIRECTOR_TECNICO", director_tecnico)
                if "AUDITOR_1" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_1", Auditor)
                if "AUDITOR_2" in cell.text.strip():
                    cell.text = cell.text.replace("AUDITOR_2", auditor_2)
                if "FECHA_AUDITORIA" in cell.text.strip():
                    cell.text = cell.text.replace("FECHA_AUDITORIA", fecha_auditoria)
                if "TIPO_DROGERIA" in cell.text.strip():
                    cell.text = cell.text.replace("TIPO_DROGERIA", tipo_drogueria)    
                    
    # Crear un string que contendrá todo el contenido a insertar en el documento
    contenido_respuestas = ""

    # Recorrer los grupos y subgrupos de preguntas
    for grupo, subgrupos in st.session_state["responses"].items():
        # Primero agregamos el nombre del grupo, que se repetirá solo una vez
        contenido_respuestas += f"**{grupo}**\n"  # Agregar el grupo en negrita
        
        for subgrupo, respuesta_data in subgrupos.items():
            # Formatear el texto para cada grupo y subgrupo
            contenido_respuestas += f"{subgrupo}\n\n"

            # Respuesta
            respuesta = respuesta_data.get('respuesta', 'sin respuesta')
            contenido_respuestas += f"**Respuesta:** {respuesta}\n"
            
            # Valor (si está presente, mostrarlo; si no, dejar en blanco)
            valor = respuesta_data.get('valor', '')
            contenido_respuestas += f"**Valor:** {valor}\n"

            # Observación (si está presente, mostrarla; si no, dejar en blanco)
            Observacion = respuesta_data.get('Observacion', '')
            contenido_respuestas += f"**Observación:** {Observacion}\n"

            # Agregar un salto de línea extra después de cada subgrupo
            contenido_respuestas += "\n"

    # Reemplazar el marcador de respuestas en el documento
    for p in doc.paragraphs:
        if "RESPUESTAS_AUDITORIA" in p.text:  # Asegúrate de tener un marcador adecuado
            p.text = p.text.replace("RESPUESTAS_AUDITORIA", contenido_respuestas)
            
            # Limpiar el texto de este párrafo antes de agregarlo
            p.clear()
            
            # Dividir el contenido en partes usando los asteriscos para las negritas
            partes = contenido_respuestas.split("**")
            for i, parte in enumerate(partes):
                # Si la parte es una palabra que debe ir en negrita (índice impar)
                if i % 2 != 0:
                    # Agregar la parte en negrita
                    run = p.add_run(parte)
                    run.bold = True
                else:
                    # Si es texto normal (índice par), añadirlo sin formato
                    p.add_run(parte)
            
    # Guardar el documento modificado
    doc.save(file_path_word)
    
    # Botón de descarga del archivo Word
    with open(file_path_word, 'rb') as word_file:
        st.download_button(label="Descargar archivo Word", data=word_file, file_name=filename_word, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    st.session_state["consecutivo"] += 1


# Función para reiniciar el formulario
def reiniciar_formulario():
    st.session_state["Datos Generales"] = {}
    st.session_state["responses"] = {}
    st.session_state["form"] = {}
    st.session_state["Operador"] = "SELECCIONAR"
    st.session_state.clear()

# Función para cargar la base de datos desde un archivo Excel
@st.cache_data
def load_data():
    url = "https://raw.githubusercontent.com/JohanaSab/proyectos/main/DIRECTORIO_Operadores.txt"
    
    response = requests.get(url)
    
    if response.status_code == 200:
        return response.content  # Retorna el contenido en binario
    
    else:
        st.error(f"❌ Error al descargar el archivo: {response.status_code}")
        return None

file_content = load_data()

if file_content is not None:
    try:
        excel_data = BytesIO(file_content)  # Convierte los bytes en un archivo en memoria
        df = pd.read_csv("DIRECTORIO_Operadores.txt", sep="|", encoding ="utf-8", on_bad_lines="skip")  # Intenta leerlo con pandas
                
    except Exception as e:
        st.error(f"⚠ Error al leer el archivo con pandas: {e}")
else:
    st.warning("No se pudo descargar el archivo.")
    
# Encabezado
st.title("")

# Barra lateral: Datos generales
st.sidebar.header("Datos Generales")
if "form" not in st.session_state:st.session_state["form"] = {}
# Seleccionar operador
Operador = st.sidebar.selectbox("Operador", Operadores.keys())

# Obtener el NIT del operador seleccionado correctamente
nit_operador = Operadores.get(Operador, "DESCONOCIDO")
st.session_state["form"]["Nit_operador"] = nit_operador
st.sidebar.write(f"NIT del Operador: {nit_operador}")

# Filtrar farmacias por NIT del operador
if nit_operador.isdigit():  # Verifica si el NIT es un número válido
    farmacias_filtradas = df[df["Nit"]== int(nit_operador)]
else:
    farmacias_filtradas = pd.DataFrame()

if not farmacias_filtradas.empty:
    # Lista desplegable con nombres de farmacias
    farmacia_seleccionada = st.sidebar.selectbox("Seleccione la farmacia:", farmacias_filtradas["NOMBRE DE LA FARMACIA"].unique())

    # Obtener datos de la farmacia seleccionada
    datos_farmacia = farmacias_filtradas[farmacias_filtradas["NOMBRE DE LA FARMACIA"] == farmacia_seleccionada].iloc[0]

    # Campos prellenados pero editables
    Nit_sucursal = st.sidebar.text_input("Nit_sucursal:", datos_farmacia["COD. SUC"])
    ciudad = st.sidebar.text_input("Ciudad:", datos_farmacia["CIUDAD / MUNICIPIO"])
    direccion = st.sidebar.text_input("Dirección:", datos_farmacia["DIRECCIÓN DE CONTACTO"])
    telefono = st.sidebar.text_input("Teléfono:", datos_farmacia["TELÉFONO DE CONTACTO"])
else:
    st.sidebar.warning("No se encontraron farmacias para este operador.")

# Otros campos editables
st.session_state["form"]["Nivel y Tipo de servicio farmacéutico"] = st.sidebar.text_input("Nivel y Tipo de servicio farmacéutico", value=st.session_state["form"].get("Nivel y Tipo de servicio farmacéutico", ""))
st.session_state["form"]["Representante legal"] = st.sidebar.text_input("Representante legal", value=st.session_state["form"].get("Representante legal", ""))
st.session_state["form"]["Director técnico"] = st.sidebar.text_input("Director técnico", value=st.session_state["form"].get("Director técnico", ""))
st.session_state["form"]["Auditor 1"] = st.sidebar.text_input("Auditor 1", value=st.session_state["form"].get("Auditor 1", ""))
st.session_state["form"]["Auditor 2"] = st.sidebar.text_input("Auditor 2", value=st.session_state["form"].get("Auditor 2", ""))
st.session_state["form"]["Fecha de Auditoria"] = st.sidebar.date_input("Fecha de Auditoria", value=pd.to_datetime(st.session_state["form"].get("Fecha de Auditoria", "2025-01-01")))
st.sidebar.radio(
    "Tipo de Droguería", ["Aliada", "Propia", "Otros"],
    key="tipo_drogueria",
    on_change=lambda: st.session_state["form"].update({"Tipo de Droguería": st.session_state["tipo_drogueria"]})
)
 
# Navegación entre grupos de preguntas
grupos = [
    "CONCEPTO NORMATIVO", "TALENTO HUMANO", "INFRAESTRUCTURA",
    "CONDICIONES LOCATIVAS", "DOTACION", "PROCESOS PRIORITARIOS",
    "GESTION DE LA CALIDAD", "APLICATIVOS TÉCNOLOGICOS (SISTEMA)"
]

grupo_actual = st.sidebar.radio("Navegación", grupos)

#Preguntas detalladas por grupo
preguntas_por_grupo = {
    "CONCEPTO NORMATIVO":["**Concepto** Cuentan con RUT con actividad comercial correspondiente al tipo de establecimiento. (Ley 1607 de 2012)\n\n**1.1** **Criterio a Verificar:** Valide que cuente con RUT vigente y describe la actividad comercial de acuerdo al tipo de establecimiento",
                          "**Concepto** N. Matricula mercantil vigente de la cámara de comercio de la respectiva jurisdicción. (Cámara de comercio renovada y actividad comercial ajustada al tipo de establecimiento). (Dec. 1879 de 2008)  \n\n**1.2** **Criterio a Verificar:** Valide que cuente con matricula mercantil vigente y que aparezca en la cámara de comercio",
                          "**Concepto** Cuentan con certificados de uso de suelo expedido por planeación municipal.  \n\n**1.3** **Criterio a Verificar:** Valide el certificado de uso de suelo de acuerdo a lo indicado",
                          "**Concepto** Cuentan con autorización de la Secretaria Seccional de Salud y Protección Social departamental o regional para su funcionamiento (Acta de visita y/o sello de la secretaria).  \n\n**1.4** **Criterio a Verificar:** Valide el acta de autorización emitida por la Secretaria Seccional de Salud y Protección Social para su funcionamiento, se debe registrar concepto y fecha de la visita, si el establecimiento no cuenta con acta debe contar con una carta de solicitud al momento de la apertura radicada al ente territorial según corresponda.",
                          "**Concepto** Cuenta con Resolución del Fondo Nacional de Estupefacientes (FNE) y/o FRE, para la tenencia y manejo de medicamentos de control especial vigente.  \n\n**1.5** **Criterio a Verificar:** Verifique la resolución del Fondo Nacional de Estupefacientes (FNE) y/o FRE y su vigencia."
                          ],
    "TALENTO HUMANO":["**Concepto** Cuenta con el talento humano necesario para la realización de las funciones de acuerdo con la constitución del establecimiento (Químico Farmacéutico, Regente de farmacia, Auxiliares técnicos en servicios farmacéuticos, etc), según sea el caso.  \n\n**2.1** **Criterio a Verificar:** El talento humano en salud, cuenta con la autorización expedida por la autoridad competente, para ejercer la profesión u ocupación.",
                      "**2.2** **Criterio a Verificar:** El personal que labora, se encuentra registrado ante las entidades departamentales o distritales de salud, y cuenta con el registro especial de prestación de salud. (RETHUS)",
                      "**2.3** **Criterio a Verificar:** Valide la existencia del manual de funciones de acuerdo al perfil establecido para los funcionarios y que se encuentre implementado.",
                      "**Concepto** En el área administrativa del servicio farmacéutico cuenta con información de los empleados.  \n\n**2.4** **Criterio a Verificar:** Validar que se cuente con el diploma ubicado en un lugar visible y que la hoja de vida se encuentre disponible en el servicio.",
                      "**Concepto** Cuenta con un cronograma de capacitaciones para los funcionarios con actas de asistencia.  \n\n**2.5** **Criterio a Verificar:** Solicite el cronograma de capacitaciones del último año, valide que se hayan realizado las capacitaciones y solicite el listado de participación. Verifique que se cuente con la adherencia del procedimiento de dispensación de medicamentos a los auxiliares de farmacia."
                      ],
    "INFRAESTRUCTURA": ["**Concepto** El servicio farmacéutico está ubicado en un lugar independiente o cualquier establecimiento comercial o habitación, con facilidad de acceso para los usuarios  \n\n**3.1** **Criterio a Verificar:** Está ubicado en área de fácil acceso y dimensiones determinadas por el volumen de actividades, el número y tipo de procesos y el número de trabajadores que laboren en el servicio.",
                        "**Concepto** Se identifica con aviso visible que exprese la razón o denominación social del establecimiento, ubicado en la parte exterior del establecimiento  \n\n**3.2** **Criterio a Verificar:** Verifique aviso al exterior del establecimiento donde se indique la razón o denominación del mismo.",
                        "**Concepto** Cuenta con unidades sanitarias correspondientes y necesarias de acuerdo al número y tipo de personas que laboran en el establecimiento.  \n\n**3.3** **Criterio a Verificar:** Verifique que cuente con las unidades sanitarias en adecuadas condiciones de mantenimiento, orden y limpieza, usada para el fin dispuesto, en proporción al número de personas que laboran en el establecimiento.",
                        "**3.4** **Criterio a Verificar:** Verifique que los baños cuenten con barandas, permitiendo la movilidad de las personas en condicion de discapacidad.",
                        "**Concepto** El tamaño de la sala de espera es acorde al volumen de la operación.  \n\n**3.5** **Criterio a Verificar:** Indague si la sala de espera del servicio farmaceutico cuenta con el metraje adecuado 20 m2"
                        ],
    "CONDICIONES LOCATIVAS": ["**Concepto** Los pisos, paredes y techos de todos los servicios deberán ser de fácil limpieza y estar en buenas condiciones de presentación y mantenimiento.  \n\n**4.1** **Criterio a Verificar:** Verifique que los pisos y paredes sean de materiales impermeables, de fácil limpieza, resistentes a factores ambientales y sistema de drenaje para su fácil limpieza y sanitización.",
                              "**4.2** **Criterio a Verificar:** Valide que los techos y cielos rasos sean resistentes, uniformes de fácil limpieza y desinfección.",
                              "**Concepto** Posee un sistema de ventilación natural y/o artificial que garantice la conservación adecuada de los medicamentos y dispositivos médicos.  \n\n**4.3** **Criterio a Verificar:** Valide que cuente con ventilación natural o artificial que garantice las condiciones adecuadas del almacenamiento de los medicamentos, recuerde que no puede haber interferencia ambiental externa dentro de los espacios (ventanas abiertas, rejillas de aire, etc)",
                              "**Concepto** Cuenta con sistema de iluminación natural o artificial para la adecuada realización de las actividades  \n\n**4.4** **Criterio a Verificar:** Valide que la iluminación sea adecuada en número e intensidad de luz para la dispensación y su correcto funcionamiento.",
                              "**Concepto** Cuenta con tomas, interruptores y cableado protegido.  \n\n**4.5** **Criterio a Verificar:** Validar que las instalaciones eléctricas como tomas, interruptores y cableado electrico y de datos esten protegidos y en buen estado.",
                              "**Concepto** Se evidencia aviso exterior e interior de: horario de atención, informativos de “Espacio libre de humo de cigarrillo”, “Requisitos de la fórmula Médica” y “Acceso restringido al personal ajeno a la farmacia”.  \n\n**4.6** **Criterio a Verificar:** Valide si se evidencia horario de atención al público.",
                              "**4.7** **Criterio a Verificar:** Verifique que a los sitios de acceso restringido solo ingrese el personal de farmacia o droguería.",
                              "**4.8** **Criterio a Verificar:** Valide si cuentan con aviso de prohibición de no fumar en el establecimiento.",
                              "**4.9** **Criterio a Verificar:** Verifique que la información este visible en cuanto a los Requisitos de la fórmula médica.",
                              "**4.10** **Criterio a Verificar:** Valide si cuentan con aviso de prohibición de consumo de alimentos y bebidas alcohólicas en el establecimiento de comercio.",
                              "**Concepto** El establecimiento tiene identificadas y señalizadas las rutas de evacuación, las salidas de emergencia y en términos generales cada una de sus áreas en caso de accidentes  \n\n**4.11** **Criterio a Verificar:** Valide que se encuentre visible el mapa de ruta de evacuación para el punto y valide que se encuentren señalizadas las rutas según el mismo.",
                              "**Concepto** El establecimiento cuenta con alertas de control de temperatura   \n\n**4.12** **Criterio a Verificar:** Revisar la temperatura mínima y máxima registrada en el equipo y solicitar justificación en caso de encontrarse desviaciones. Para casos en los que se cuente con sensor netux abra la nevera y espere a que se encuentre a más de 8°C, valide que se active la alerta de temperatura fuera de rango.",
                              "**Concepto** Validar que la temperatura en el momento de la auditoría se encuentran dentro de los rangos establecidos.   \n\n**4.13** **Criterio a Verificar:** Qué temperatura registra el termohigrómetro actualmente? ____°C Qué % de humedad registra el termohigrómetro actualmente? _____%.",
                              "**4.14** **Criterio a Verificar:** Alguno de esas temperaturas registradas es mayor a 30°C? Si la respuesta es SI anotar cual o cuales fechas registró esa temperatura mayor a 30°C en la ultima casilla de Observaciones / Hallazgos. (Si la respuesta fue Si Calificar con 0 en caso contrario calificar con 100).",
                              "**4.15** **Criterio a Verificar:** Alguno de esos % de humedad registradas es mayor a 75%? Si la respuesta es SI anotar cual o cuales fechas registró esa humedad mayor a 75% en la última casilla de Observaciones / Hallazgos.",
                              "**Concepto** En las áreas de almacenamiento de medicamentos y  dispositivos  médicos  cuenta con alarmas  sensibles  al  humo  y extintores de incendios. En éstas no se podrán acumular residuos.  \n\n**4.16** **Criterio a Verificar:** Verifique la disposición de alarmas sensibles al humo y extintores, además verifique que cuentan con extintores en funcionamiento, ubicados de acuerdo a las normas de seguridad, con fecha vigente de recarga y señalizada la ubicación.",
                              "**Concepto** Cuenta con las áreas requeridas para la realizacion de los procesos del establecimiento. Diferentes para cada proceso, diferenciadas y debidamente señalizadas.  \n\n**4.17** **Criterio a Verificar:** Verifique que cuenta con las áreas de- Área administrativa debidamente delimitada o señalizada con rótulo de identificación.",
                              "**4.18** **Criterio a Verificar:** Área para la recepción de productos: Farmacéuticos, dispositivos médicos y productos autorizados, debidamente delimitada o señalizada con rótulo de identificación.",
                              "**4.19** **Criterio a Verificar:** Área de cuarentena de medicamentos debidamente señalizada, en ella también se podrán almacenar de manera transitoria los productos retirados del mercado con rótulo de identificación.",
                              "**4.20** **Criterio a Verificar:** Área del almacenamiento de medicamentos y dispositivos médicos independiente, diferenciada y señalizada con rótulo de identificación.",
                              "**4.21** **Criterio a Verificar:** Área independiente y segura para el almacenamiento de medicamentos de control especial con rótulo de identificación y con acceso restringido al publico general.",
                              "**4.22** **Criterio a Verificar:** Área especial, debidamente identificada, para el almacenamiento transitorio de los medicamentos vencidos o deteriorados que deban ser técnicamente destruidos o desnaturalizados con rótulo de identificación (área especial para devoluciones).",
                              "**4.23** **Criterio a Verificar:** Área de dispensación  de medicamentos y entrega de dispositivos médicos y productos autorizados, debidamente delimitada o señalizada con rótulo de identificación.",
                              "**4.24** **Criterio a Verificar:** Área para el manejo y disposición de residuos con rótulo de identificación.",
                              "**4.25** **Criterio a Verificar:** Área destinada para el almacenamiento de productos rechazados, devueltos y retirados del mercado con rótulo de identificación (área especial para devoluciones).",
                              "**4.26** **Criterio a Verificar:** Área para el almacenamiento de medicamentos fitoterapéuticos, cuando estos se vendan al detal al público con rótulo de identificación.",
                              "**Concepto** Las instalaciones se mantienen en condiciones adecuadas de conservación higiene y limpieza.  \n\n**4.27** **Criterio a Verificar:** Verifica que todas las áreas se encuentran limpias.",
                              "**Concepto** Área de aseo, donde se encuentre los implementos para la adecuada higienización.  \n\n**4.28** **Criterio a Verificar:** Verifique que los desechos son clasificados en su fuente de generación según el código de colores: Blanca, negra y verde.",
                              "**Concepto** Lista de chequeo con frecuencias de desinfección y limpieza de áreas.  \n\n**4.29** **Criterio a Verificar:** Valide que la lista de chequeo se encuentre completamente diligenciada con frecuencias de desinfección y limpieza de áreas.",
                              "**Concepto** Manejo y disposicion de residuos.  \n\n**4.30** **Criterio a Verificar:** El establecimiento cuenta con el plan de gestión integral de residuos sólidos (ordinarios).",
                              "**4.31** **Criterio a Verificar:** Se cuenta con recipientes suficientes, con tapa y bolsa de colores para clasificación de residuos, bien ubicados e identificados para recolección interna de los desechos según corresponda.",
                              "**4.32** **Criterio a Verificar:** Cuenta con un contrato o acuerdo comercial de recolección con empresa responsable de la recolección de residuos.",
                              "**4.33** **Criterio a Verificar:** Cuenta con elementos de aseo para el área administrativa y para el área de farmacia.",
                              "**4.34** **Criterio a Verificar:** Retiro de basuras con frecuencia necesaria y la ruta."
                              ],
    "DOTACION": ["**Concepto** Cuenta con la dotación y muebles exclusivos y necesarios para la adquisición, recepción, almacenamiento, conservación (como manejo de cadena de frio, medicamentos fotosensibles, higroscópicos entre otros) y dispensación de los medicamentos y dispositivos médicos para la realización de los procesos que ofrezcan, de acuerdo con las recomendaciones dadas por los fabricantes, además de las condiciones de temperatura y humedad, registro de condiciones ambientales, termohigrometros calibrados.  \n\n**5.1** **Criterio a Verificar:** Verifique que cuenta con estantería de material sanitario, impermeable y fácil de limpiar. (estibas plasticas, no de madera) suficiente para el almacenamiento de los medicamentos de acuerdo a su condición y cantidad.",
                 "**5.2** **Criterio a Verificar:** Cuenta con equipos para la medición de condiciones de temperatura y humedad y los mismos se encuentran en funcionamiento y calibrados. Solicite las hojas de vida y las certificados de calibración. Valide que la empresa de calibración se encuentre certificada por la ONAC. Solicite el cronograma de mantenimiento y certificados de los equipos refrigerantes.",
                 "**5.3** **Criterio a Verificar:** Solicite los registros de condiciones de temperatura y humedad, verifique los datos actualizados y que se encuentren en los niveles de seguridad establecidos, en caso de tener algún dato fuera del rango solicite la acción correctiva realizada para el caso.",
                 "**5.4** **Criterio a Verificar:** Que temperatura registra el termómetro o termohigrómetro de la nevera actualmente? ____°C.",
                 "**5.5** **Criterio a Verificar:** Alguna de esas temperaturas registradas es menor a 2°C o mayor a 8°C? Si la respuesta es SI anotar cual fecha registró esas temperaturas en la ultima casilla (Observaciones) (Si la respuesta fue Si Calificar con 0 en caso contrario calificar con 100).",
                 "**Concepto** Se evidencia equipos e implementos de seguridad en funcionamiento y ubicados donde tengan fácil acceso (extintores, barandas, estibas, extractoras).  \n\n**5.6** **Criterio a Verificar:** Cuentan con extintores en funcionamiento, ubicados de acuerdo a las normas de seguridad, con fecha vigente de recarga y señalizada su ubicación.",
                 "**Concepto** Cuenta con los equipos necesarios para cumplir con el plan de contingencia, en caso de falla eléctrica para mantener la cadena de frío.   \n\n**5.7** **Criterio a Verificar:** Valide el plan de contingencia que se tiene implementado en caso de falla de temperaura y verifique que se cuente con los insumos necesarios para esta acción.",
                 "**Concepto** Cuenta con medios refrigerantes para el almacenamiento de medicamentos que lo requieren.  \n\n**5.8** **Criterio a Verificar:** Verifique que las neveras sean de uso exclusivo para almacenamiento de los medicamentos que lo requieren.",
                 "**Concepto** Cuenta con literatura científica disponible y aceptada internacionalmente; referencias  bibliográficas actualizadas y confiables sobre farmacología y farmacoterapia para su consulta.  \n\n**5.9** **Criterio a Verificar:** Indage si tiene fuentes de cosulta frente temas relacionados con farmacovigilancia o dudas relacionadas con usos de medicamentos.",
                 "**Concepto** Cuenta con Kit de manejo de derrames y ruptura de medicamentos.  \n\n**5.10** **Criterio a Verificar:** Verifique que este kit se encuentre completo para su uso."
                 ],
    "PROCESOS PRIORITARIOS": ["**Concepto** Recepción  \n\n**6.1** **Criterio a Verificar:** Solicite el documento del proceso de recepción técnica administrativa de medicamentos, indague como se realiza y valide en las actas su correcta ejecución.",
                              "**6.2** **Criterio a Verificar:** En caso que transporten medicamentos y dispositivos médicos a otras sedes o sucursales, verifican las condiciones en que son transportados y llevan registro de dicho procedimiento (documento soporte de traslado).",
                              "**6.3** **Criterio a Verificar:** Se tienen en cuenta las condiciones dadas por el fabricante para el almacenamiento de medicamentos y dispositivos médicos.",
                              "**6.4** **Criterio a Verificar:** Verificar el procedimiento de recepción de medicamentos y dispositivos médicos que llegan del CEDI o bodega central a las droguerías, verificando mantenimiento de la cadena de frio y condiciones de embalaje.",
                              "**Concepto** Almacenamiento  \n\n**6.5** **Criterio a Verificar:** Solicite el documento del proceso de almacenamiento y valide que se realice en la manera descrita, tener en cuenta que debe ser acuerdo a un método de clasificación (alfabeticamente, forma farmacéutica, otros) y los dispositivos médicos deben almacenarse de acuerdo a lo establecido, siempre y cuando garantice el orden y minimice eventos de confusión.",
                              "**6.6** **Criterio a Verificar:** Se realiza la identificación y marcación de medicamentos lasa y alto riesgo.",
                              "**6.7** **Criterio a Verificar:** El establecimiento está libre de productos farmacéuticos prohibidos (medicamentos de control especial sin autorización, vencidos en el área de dispensación, muestras médicas, productos con enmendaduras o etiquetas que oculten información, sin registro sanitario, con el sistema de seguridad alterado, entre otros).",
                              "**6.8** **Criterio a Verificar:** Solicite el proceso de control de fechas de vencimiento de medicamentos y valide que se lleve a cabo y que los auxiliares lo conozcan. Adicionalmente valide que no haya medicamentos vencidos ubicados en las estanterías de suministro.",
                              "**Concepto** Medicamentos de control especial  \n\n**6.9** **Criterio a Verificar:** Verifique que el área para el almacenamiento de los medicamentos de control especial sea independiente, diferenciada, señalizada y se encuentra bajo llave en área de acceso restringido al público general. Solo el director técnico debe tener acceso, debe solicitarse el soporte que delega la responsabilidad a los auxiliares para cuando no este el director técnico y debe especificar nombre y cédula del responsable.",
                              "**6.10** **Criterio a Verificar:** Cumple con la presentación de informe al FNE y/o FRE dentro de las fechas establecidas. (Validar en ciudades fuera de Bogotá y Cundinamarca estas se validaran a nivel central).",
                              "**6.11** **Criterio a Verificar:** Tome al menos 10 medicamentos aleatoriamente y valide las existencias Vs documento de control de dispensación (libro). Debe coincidir la existencia física Vs lo reportado en el libro.",
                              "**Concepto** Dispensación  \n\n**6.12** **Criterio a Verificar:** Solicite el documento con el proceso  de dispensación de medicamentos y dispositivos medicos y otros productos farmaceuticos, valide que se esté ejecutando de acuerdo a lo descrito.",
                              "**6.13** **Criterio a Verificar:** Valide que se esté colocando el sello de dispensado en los medicamentos de control especial",
                              "**6.14** **Criterio a Verificar:** Verifique aleatoriamente si al paciente no se le realiza entrega completa de la orden medica por algún medicamento que no cuenta con novedad en inventario, que a este se le genere el pendiente en sistema y se le ofrezca el envío a domicilio del mismo. Llevar consigo la última versión del informe de novedades.",
                              "**Concepto** Atención farmacéutica  \n\n**6.15** **Criterio a Verificar:** Cuentan con Químico Farmacéutico en caso de que se requiera la asesoria farmacológica.",
                              "**Concepto** Errores de medicación, Farmacovigilancia y Tecnovigilancia  \n\n**6.16** **Criterio a Verificar:** Los funcionarios de la droguería conoce los programas de farmacovigilancia, tecnovigilancia institucionales y errores de medicación.",
                              "**Concepto** Alertas sanitarias  \n\n**6.17** **Criterio a Verificar:** Verifique si se realiza la socialización de las alertas Sanitarias al personal de la farmacia emitidas por los entes regulatorios."
                              ],
    "GESTION DE LA CALIDAD": ["**Concepto** Cuenta con visitas de entes de control con algún tipo de hallazgos y los mismos ya se encuentran solventados  \n\n**7.1** **Criterio a Verificar:** Valide con soporte físico la útlima visita de entes de control y verifique si hay requerimientos y los mismos ya se encuentran solventados.",
                             "**Concepto** Realizan inventario (periódico, trimestral o semestral)  \n\n**7.2** **Criterio a Verificar:** Solicite el inventario con su última actualización.",
                             "**Concepto** Establecimiento cuenta con matriz de riesgo y la misma es de conocimiento por parte de los funcionarios.  \n\n**7.3** **Criterio a Verificar:** Solicitar la matriz de riesgo y validar que los funcionarios la conozcan.",
                             "**Concepto** Contar cajas o ventanilla de atención.  \n\n**7.4** **Criterio a Verificar:** ¿Las cajas son suficientes para la atención de los usuarios? (Volumen de personas Vs. tiempo de atención).",
                             "**7.4.1** **Criterio a Verificar:** ¿Cuantas cajas se encuentran habilitadas?",
                             "**Concepto** Validar si se tiene digiturno o turnos de atención.  \n\n**7.5** **Criterio a Verificar:** ¿Se cuenta con digiturno?",
                             "**7.7** **Criterio a Verificar:** ¿Cuentan con la priorización en el digiturno: General y Preferencial?",
                             "**Concepto** Si tiene digiturno y en el momento de la auditoria hay pacientes para ser atendidos, tomar un turno, y registrar cuanto tiempo demoró desde la toma del turno hasta que fue llamado en ventanilla.   \n\n**7.8** **Criterio a Verificar:** Cuanto tiempo demoró el turno en ser llamado a la ventanilla? ______." 
                             ],
    "APLICATIVOS TÉCNOLOGICOS (SISTEMA)": ["**Concepto** Cuenta con medios, preferiblemente computarizados que permitan el registro de adecuado de formulas prescritas, medicamentos de control especial y demas productos que lo requieran.  \n\n**8.1** **Criterio a Verificar:** Verificar el sistema que manejan en el punto, su correcto funcionamiento y utilización.",
                                           "**8.2** **Criterio a Verificar:** Validar de manera aleatoria en el informe de novedades que los medicamentos que presentan novedad (desabastecidos) se encuentran bloqueados en el aplicativo de dispensación. Validar para los casos en los que no cuenta con disponibilidad en el punto la generación del pendiente en el sistema."
                                           ]
    }

# Mostrar preguntas del grupo actual
st.header(grupo_actual)
preguntas = preguntas_por_grupo.get(grupo_actual,[])

# Crear estructura de almacenamiento si no existe
if grupo_actual not in st.session_state["responses"]:
    st.session_state["responses"][grupo_actual] = {}

for pregunta in preguntas:
    if pregunta not in st.session_state["responses"][grupo_actual]:
        st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": None,
            "valor": None,
            "Observacion": "",
        }
        
       
        
    #respuestas seleccionadas previamente
    respuesta_actual = st.session_state["responses"][grupo_actual][pregunta].get("respuesta","Seleccione una opcion")
    observacion_actual = st.session_state["responses"][grupo_actual][pregunta].get("Observacion","")

    #si no hay respuesta mostrar una opcion vacia para obligar la seleccion
    Opciones = ["Selecciona una opción","Cumple totalmente", "Cumple parcialmente", "Incumple totalmente", "No Aplica"]

    # Mostrar la pregunta y permitir selección
    respuesta = st.radio(
        pregunta,
        Opciones,
        key=f"respuesta_{grupo_actual}_{pregunta}",
        index = Opciones.index(respuesta_actual) if respuesta_actual in Opciones else 0
        )
    
    #Campo de observacion
    Observacion = st.text_area(f"Observacion",
                               value=observacion_actual,
                               key=f"Observacion_{grupo_actual}_{pregunta}",
                               on_change=lambda p=pregunta: actualizar_observacion(grupo_actual,p)
                               )
    def actualizar_observacion(grupo,pregunta):
        """Actualizar la observacion en session_state al escribir en el campo de texto."""
        clave = f"Observacion_{grupo}_{pregunta}"
        st.session_state["responses"][grupo][pregunta]["Observacion"] = st.session_state[clave]
        
    # Guardar respuesta y valor
    if respuesta != "Selecciona una opción":
            valor = 100 if respuesta == "Cumple totalmente" else 50 if respuesta == "Cumple parcialmente" else 0
            if grupo_actual not in st.session_state["responses"]:
                st.session_state["responses"][grupo_actual] = {}
            st.session_state["responses"][grupo_actual][pregunta] = {
            "respuesta": respuesta,
            "valor": valor,
            "Observacion": st.session_state[f"Observacion_{grupo_actual}_{pregunta}"]  # Usar el estado guardado
    }
   

# Calcular promedio
    total_valores = sum(
        r["valor"] if r ["valor"] is not None else 0 for g in st.session_state["responses"].values() for r in g.values()
)
num_preguntas = sum(len(g) for g in st.session_state["responses"].values())

#asegurar que el promedio sea sobre 100
promedio = (total_valores / (num_preguntas*100)) *100 if num_preguntas > 0 else 0

st.sidebar.metric("Promedio Total", round(promedio,2))

# Botones
st.sidebar.button("Finalizar y Enviar", on_click=finalizar_formulario)
st.sidebar.button("Nuevo formulario", on_click=reiniciar_formulario)
st.sidebar.button("Guardar", on_click=guardar_estado)

# Entrada de usuario
consecutivo_input = st.sidebar.text_input("Ingrese el código del formulario a cargar:")

# Botón para cargar el formulario
if st.sidebar.button("Cargar formulario"):
    if consecutivo_input:
        # Indicar que se debe cargar el formulario
        st.session_state["load_form"] = True
        st.session_state["input_consecutivo"] = consecutivo_input
    else:
        st.error("Por favor, ingrese un código de formulario válido.")

# Cargar el formulario si se indicó
if st.session_state.get("load_form"):
    consecutivo_input = st.session_state.get("input_consecutivo")
    cargar_formulario_por_consecutivo(consecutivo_input)  # Llamada correcta pasando el argumento
    st.session_state["load_form"] = False  # Restablecer el indicador
